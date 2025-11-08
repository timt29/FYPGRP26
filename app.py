from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, Response, abort, make_response,g
import os, io, re, time, uuid, time, zipfile, shutil, secrets
from functools import wraps
import webbrowser
import threading
import mysql.connector
import json
import decimal
import traceback
import re
from datetime import datetime, timezone
from werkzeug.utils import secure_filename
from rapidfuzz import process, fuzz
import fitz
from pathlib import Path
from docx import Document
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from deep_translator import GoogleTranslator
from utils.profanity_filter import contains_profanity, censor_text
from utils.content_analyzer import analyze_content
from typing import Optional
from utils.chatbot import getChatbotResponse
import psycopg2
from psycopg2.extras import RealDictCursor
import nltk
import socket

# Path to bundled nltk_data (inside your repo)
nltk_data_dir = os.path.join(os.path.dirname(__file__), "nltk_data")
nltk.data.path.append(nltk_data_dir)
# Test
from nltk.tokenize import sent_tokenize
print(sent_tokenize("Hello world. This is a test."))


app = Flask(__name__)
app.secret_key = "your_secret_key"  # Replace with a secure key

'''
# ---------- DB ----------
def get_db_connection():
    return mysql.connector.connect(
        host="localhost",
        user="root",
        password="password",
        database="nrs"
    )
'''

def get_db_connection():
    database = os.getenv("database", "postgres")
    user = os.getenv("user", "postgres.ioxwfemawqhuqwkmshoc")
    password = os.getenv("password", "ZavXqAKfvdBvKRUKrjhQZoMYypyHLQes")

    pooled_host = "aws-1-ap-southeast-1.pooler.supabase.com"
    direct_host = "db.ioxwfemawqhuqwkmshoc.supabase.co"

    try:
        pooled_ipv4 = socket.gethostbyname(pooled_host)
        print(f"Connecting to pooled IPv4: {pooled_ipv4}:6543")

        conn = psycopg2.connect(
            host=pooled_ipv4,
            port=6543,
            database=database,
            user=user,
            password=password,
            sslmode="require"
        )
        print("✅ Connected to Supabase (pooled IPv4)")
        return conn

    except psycopg2.OperationalError as e:
        print(f"❌ Pooled connection failed: {e}")

        # fallback to direct
        try:
            direct_ipv4 = socket.gethostbyname(direct_host)
            print(f"Connecting to direct IPv4: {direct_ipv4}:5432")

            conn = psycopg2.connect(
                host=direct_ipv4,
                port=5432,
                database=database,
                user=user,
                password=password,
                sslmode="require"
            )
            print("✅ Connected to Supabase (direct IPv4)")
            return conn

        except psycopg2.OperationalError as e2:
            print(f"❌ Direct connection failed: {e2}")
            raise ConnectionError("Could not connect to Supabase via pooled or direct IPv4")


print("Connected to MySQL")

def get_cursor(conn):
    return conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)

# ---------- Auth wrapper ---------- # (YY)
from functools import wraps
from flask import session, redirect, url_for, flash

def login_required(*roles):
    """
    Usage:
      @login_required()                        -> requires login (any role)
      @login_required("Subscriber")            -> only Subscriber
      @login_required("Author","Subscriber")   -> either role
      @login_required(("Author","Subscriber")) -> tuple/list also works
    """
    # Normalize to a set; None means "any role"
    if len(roles) == 1 and isinstance(roles[0], (list, tuple, set)):
        allowed = set(roles[0])
    elif roles:
        allowed = set(roles)
    else:
        allowed = None

    def wrapper(fn):
        @wraps(fn)
        def decorated_view(*args, **kwargs):
            # If no session → force logout
            if "userid" not in session:
                session.clear()
                return redirect(url_for("login"))

            # role check if specific roles required
            if allowed is not None and session.get("usertype") not in allowed:
                session.clear()
                flash("Access denied. Please log in again.", "danger")
                return redirect(url_for("login"))

            return fn(*args, **kwargs)
        return decorated_view
    return wrapper

# ---------- Public pages ----------# (TIM)
@app.route("/")
def index():
    conn = get_db_connection()
    cursor = get_cursor(conn)
    cursor.execute("""
        SELECT articleid, title, content, author, published_at, updated_at, image
        FROM articles 
        WHERE visible = 1
        ORDER BY published_at DESC
        LIMIT 10
    """)
    articles = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template("index.html", articles=articles)

@app.route("/article/<int:article_id>")
def article(article_id):
    conn = get_db_connection()
    cursor = get_cursor(conn)
    cursor.execute("SELECT * FROM articles WHERE articleID = %s", (article_id,))
    article = cursor.fetchone()
    cursor.close()
    conn.close()

    #if not article: # i added this part (ZT)
        #return "Article not found", 404
    
    # --- AI Summarizer call ---
    summary_points = []
    try:
        full_text = article.get("content", "").strip()
        summary_points = summarize_to_points(full_text)
    except Exception as e:
        summary_points = [f"AI summarizer not available: {str(e)}"]

    return render_template("article.html", article=article, summary_points=summary_points)

#(MW) ----- Increment views (+1 once per session per article)
@app.route("/api/article/<int:article_id>/view", methods=["POST"])
def api_article_view_beacon(article_id: int):

    # Track which article IDs this session has viewed
    viewed_key = "viewed_articles"
    try:
        already = set(int(x) for x in session.get(viewed_key, []))
    except Exception:
        already = set()

    conn = get_db_connection()
    cur  = get_cursor(conn)

    updated = False
    if article_id not in already:
        cur.execute("UPDATE articles SET views = views + 1 WHERE articleid = %s", (article_id,))
        conn.commit()
        already.add(article_id)
        session[viewed_key] = list(already)
        updated = True

    # (Optional) return up-to-date views
    cur.execute("SELECT COALESCE(views,0) AS views FROM articles WHERE articleid = %s", (article_id,))
    row = cur.fetchone() or {"views": 0}
    cur.close(); conn.close()

    return jsonify({"ok": True, "updated": updated, "views": int(row["views"])})

@app.route("/search")
def search():
    query = request.args.get("q", "").strip()
    no_results = False
    articles = []
    fallback_articles = []

    conn = get_db_connection()
    cursor = get_cursor(conn)

    # --- 1. Perform the search ---
    if query:
        cursor.execute("""
            SELECT articleid, title, content, author, published_at, updated_at, image
            FROM articles
            WHERE title LIKE %s OR content LIKE %s
            ORDER BY published_at DESC
        """, (f"%{query}%", f"%{query}%"))
        articles = cursor.fetchall()

    # --- 2. If no results, show recent articles as fallback ---
    if not articles:
        no_results = True
        cursor.execute("""
            SELECT a.articleid, a.title, a.content, a.author, a.published_at, a.image, c.name AS category
            FROM articles a
            JOIN categories c ON a.catid = c.categoryid
            ORDER BY a.published_at DESC
            LIMIT 10
        """)
        fallback_articles = cursor.fetchall()

    cursor.close()
    conn.close()

    # --- 3. Render template with both lists ---
    return render_template(
        "index.html",
        articles=articles,
        fallback_articles=fallback_articles,
        no_results=no_results,
        search_query=query
    )

def get_categories():
    conn = get_db_connection()
    cursor = get_cursor(conn)
    cursor.execute("SELECT name FROM categories")
    categories = cursor.fetchall()
    conn.close()
    return categories

@app.context_processor
def inject_categories():
    return dict(nav_categories=get_categories())

@app.route("/category/<category_name>")
def category(category_name):
    conn = get_db_connection()
    cursor = get_cursor(conn)

    cursor.execute("""
        SELECT a.articleid, a.title, a.content, a.author, a.published_at, a.image, c.name AS category
        FROM articles a
        JOIN categories c ON a.catid = c.categoryid
        WHERE c.name = %s AND a.visible = 1
        ORDER BY a.published_at DESC
    """, (category_name,))
    articles = cursor.fetchall()

    conn.close()

    # Pass a flag if no articles found
    return render_template("index.html", articles=articles, category_name=category_name, no_articles=(len(articles) == 0))

# ---------- Role homepages ---------- # (YY)
@app.route("/privacy")
def privacy():
    return render_template("privacy.html")

@app.route("/terms")
def terms():
    return render_template("terms.html")

@app.route("/adminHomepage")
@login_required("Admin")
def adminHomepage():
    conn = get_db_connection()
    cursor = get_cursor(conn)

    # Fetch user stats
    cursor.execute("SELECT COUNT(*) AS total FROM users")
    total_users = cursor.fetchone()["total"]

    cursor.execute("""
    SELECT COUNT(*) AS new_today
    FROM users
    WHERE created_at::date = CURRENT_DATE
""")
    new_users_today = cursor.fetchone()["new_today"]

    cursor.execute("SELECT COUNT(*) AS suspended FROM users WHERE usertype='Suspended'")
    suspended_users = cursor.fetchone()["suspended"]

    cursor.execute("SELECT COUNT(*) AS active FROM users WHERE usertype!='Suspended'")
    active_users = cursor.fetchone()["active"]

    cursor.close()
    conn.close()

    # Render template directly; @after_request will handle no-cache headers
    return render_template(
        "adminHomepage.html",
        total_users=total_users,
        new_users_today=new_users_today,
        suspended_users=suspended_users,
        active_users=active_users,
    )

# (YY)
@app.route("/modHomepage")
@login_required("Moderator")
def modHomepage():
    conn = get_db_connection()
    cursor = get_cursor(conn)

    try:
        # Count flagged articles with pending status
        cursor.execute("""
            SELECT COUNT(*) AS total
            FROM article_reports
            WHERE status = 'pending'
        """)
        flagged_articles = cursor.fetchone()["total"]

        # Count flagged comments with pending status
        cursor.execute("""SELECT COUNT(*) AS total FROM comment_reports WHERE status = 'pending' """)
        flagged_comments = cursor.fetchone()["total"]

        # Optionally: count pending articles
        cursor.execute("SELECT COUNT(*) AS total FROM articles WHERE status = 'pending'")
        pending_articles = cursor.fetchone()["total"]

    except mysql.connector.Error as err:
        print("Database error:", err)
        flagged_articles = flagged_comments = pending_articles = 0

    finally:
        cursor.close()
        conn.close()

    return render_template(
        "modHomepage.html",
        flagged_articles=flagged_articles,
        flagged_comments=flagged_comments,
        pending_articles=pending_articles
    )

# (YY)
@app.route("/subscriberHomepage")
@login_required("Subscriber" , "Author")
def subscriberHomepage():
    conn = get_db_connection()
    cur = get_cursor(conn)
    # (existing warnings query)
    cur.execute("SELECT message, created_at FROM warnings WHERE userid=%s", (session["userid"],))
    warnings = cur.fetchall()
    # Pull categories dynamically
    cur.execute("SELECT categoryid, name FROM categories ORDER BY categoryid ASC")
    categories = cur.fetchall()
    cur.close(); conn.close()
    # active tab via query param (?cat=Technology), default 'Trending'
    active_category = request.args.get("cat", "Trending")
    return render_template(
        "subscriberHomepage.html",
        warnings=warnings,
        categories=categories,
        active_category=active_category
    )
	
# (MW)
# API: articles feed for subscriberHomepage
@app.route("/api/articles")
@login_required("Author", "Subscriber")
def api_articles():
    cat    = request.args.get("cat")
    q      = (request.args.get("q") or "").strip()
    limit  = int(request.args.get("limit", 10))
    offset = int(request.args.get("offset", 0))

    conn = get_db_connection()
    cur = get_cursor(conn)

    params = []
    where  = ["a.draft = 0", "a.visible = 1"]   # only published & visible

    # category filter
    join_cat = "LEFT JOIN categories c ON a.catid = c.categoryid"
    if cat:
        where.append("c.name = %s")
        params.append(cat)

    # search filter
    if q:
        where.append("(a.title LIKE %s OR a.content LIKE %s)")
        like = f"%{q}%"
        params.extend([like, like])

    sql = f"""
        SELECT a.articleid, a.title, a.content, a.author,
               a.published_at, a.updated_at, a.image, c.name AS category
        FROM articles a
        {join_cat}
        WHERE {" AND ".join(where)}
        ORDER BY
          CASE WHEN (EXISTS (
            SELECT 1 FROM users u
            WHERE TRIM(LOWER(u.name)) = TRIM(LOWER(a.author))
              AND TRIM(LOWER(u.usertype)) = 'author'
          )) THEN 0 ELSE 1 END,
          COALESCE(a.published_at, a.updated_at) DESC,
          a.articleid DESC
        LIMIT %s OFFSET %s
    """
    params.extend([limit, offset])

    cur.execute(sql, tuple(params))
    rows = cur.fetchall()
    cur.close(); conn.close()
    return jsonify(rows)

@app.route("/subscriber/create-article")
@login_required("Author", "Subscriber")
def subscriberCreateArticle():
    return render_template("subscriberCreateArticle.html")

@app.route("/api/categories")
@login_required("Author", "Subscriber")
def api_categories():
    conn = get_db_connection()
    cur = get_cursor(conn)
    cur.execute("SELECT categoryid, name FROM categories ORDER BY categoryid ASC")
    rows = cur.fetchall()
    cur.close(); conn.close()
    return jsonify(rows)

# --- Profanity filter --- # (NICOLE)
def check_profanity(title, content):
    try:
        if contains_profanity(title) or contains_profanity(content):
            return {
                "ok": False,
                "message": "Your article contains blocked words. Please revise and try again."
            }
    except Exception as e:
        print("⚠️ Profanity filter error:", e)
        # Don’t block upload if profanity detection fails
        return {"ok": True, "message": None}

    return {"ok": True, "message": None}

# --- AI moderation --- # (NICOLE)
def moderate_article(cur, title, content):
    publish = True
    visible = 1
    flash_msg = "Article uploaded successfully."

    try:
        combined_text = f"{title}\n{content}"
        moderation_result = analyze_content(combined_text)

        if moderation_result['status'] in ("BLOCKED", "FLAGGED"):
            publish = False
            visible = 0
            flash_msg = "Article sent for moderator review due to content issues."

            # --- AI auto-report if flagged or blocked ---
            try:
                article_id = cur.lastrowid
                cur.execute("""
                    INSERT INTO article_reports (article_id, reporter_id, reason, details, status, created_at, updated_at)
                    VALUES (%s, %s, %s, %s, 'pending', NOW(), NOW())
                """, (
                    article_id,
                    session.get("userid") or 0,
                    "AI moderation",
                    moderation_result.get("reason", "")
                ))
                print(f"✅ AI auto-report created for article {article_id}")
            except Exception as e:
                print("⚠️ AI report insert failed:", e)

    except Exception as e:
        print("AI moderation error:", e)

    return publish, visible, flash_msg

import os, time, json, traceback, shutil
from datetime import datetime
from werkzeug.utils import secure_filename

MAX_IMPORT_IMAGES = 3
ALLOWED_IMAGE_EXTS = {'.png', '.jpg', '.jpeg', '.gif', '.webp'}

def _copy_static_file(rel_path, target_dir):
    if not rel_path or not rel_path.startswith("/static/"):
        return None
    src_path = os.path.join(app.root_path, rel_path.lstrip("/"))
    if not os.path.exists(src_path):
        return None
    os.makedirs(target_dir, exist_ok=True)
    name, ext = os.path.splitext(os.path.basename(src_path))
    unique = f"{name}_{int(time.time())}{ext}"
    dst_path = os.path.join(target_dir, unique)
    shutil.copy2(src_path, dst_path)
    return f"/static/img/{unique}"

@app.route("/api/articles", methods=["POST"])
@login_required("Author", "Subscriber")
def api_articles_create():
    try:
        user = session.get("user")
        if not user:
            return jsonify({"ok": False, "message": "User not logged in"}), 401

        title = (request.form.get("title") or "").strip()
        content = (request.form.get("content") or "").strip()
        category_id = (request.form.get("category_id") or "").strip()
        action = (request.form.get("action") or "publish").lower()
        publish = (action == "publish")

        if not title or not category_id.isdigit():
            return jsonify({"ok": False, "message": "Title and category required"}), 400

        # handle uploaded image
        image_rel = None
        img_dir = os.path.join(app.root_path, "static", "img")
        os.makedirs(img_dir, exist_ok=True)
        file_up = request.files.get("image")
        if file_up and file_up.filename:
            ext = os.path.splitext(file_up.filename)[1].lower()
            if ext in ALLOWED_IMAGE_EXTS:
                safe_name = secure_filename(file_up.filename)
                unique = f"{safe_name}_{int(time.time())}{ext}"
                file_up.save(os.path.join(img_dir, unique))
                image_rel = f"/static/img/{unique}"

        # handle import images (first 3 only)
        try:
            import_images = json.loads(request.form.get("import_images", "[]"))
        except Exception:
            import_images = []
        for rel in import_images[:MAX_IMPORT_IMAGES]:
            copied = _copy_static_file(rel, img_dir)
            if copied and not image_rel:
                image_rel = copied

        # DB insert (draft or publish)
        conn = get_db_connection()
        cur = get_cursor(conn)
        cur.execute("""
            INSERT INTO articles
            (title, content, author, published_at, updated_at, image, catid, draft, visible, status)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, (
            title,
            content,
            user,
            datetime.now() if publish else None,
            datetime.now(),
            image_rel,
            category_id,
            0 if publish else 1,
            1,
            'published' if publish else 'draft'
        ))
        conn.commit()
        cur.close()
        conn.close()

        return jsonify({
            "ok": True,
            "message": "Draft saved." if not publish else "Article published.",
            "redirect": url_for("subscriberHomepage")
        }), 200

    except Exception:
        print(traceback.format_exc())
        return jsonify({"ok": False, "message": "Internal server error"}), 500




# (MW) --------Import document, extract text & images--------
@app.post("/subscriber/api/import-article")
@login_required("Author", "Subscriber")
def api_import_article():
    file = request.files.get("file")
    if not file or not file.filename:
        return jsonify(ok=False, message="No file provided"), 400

    ext = os.path.splitext(file.filename.lower())[1]
    if ext not in {".txt", ".pdf", ".doc", ".docx"}:
        return jsonify(ok=False, message="Unsupported file type"), 400

    raw = file.read()
    if not raw:
        return jsonify(ok=False, message="Empty file"), 400

    base_id = uuid.uuid4().hex[:12]
    base_dir = os.path.join(app.static_folder, "uploads", "articles", base_id)
    img_dir  = os.path.join(base_dir, "img")
    os.makedirs(img_dir, exist_ok=True)

    text = ""
    img_names = []

    try:
        if ext == ".txt":
            try:
                text = raw.decode("utf-8")
            except UnicodeDecodeError:
                text = raw.decode("latin-1", errors="ignore")
        elif ext == ".pdf":
            text, img_names = _extract_from_pdf(raw, img_dir, base_id)
        elif ext == ".docx":
            text = _extract_text_docx(raw)
            img_names = _extract_images_docx(raw, img_dir, base_id)
        elif ext == ".doc":
            try:
                docx_bytes = _convert_doc_to_docx(raw)
            except Exception as e:
                return jsonify(ok=False, message=f".doc conversion failed: {e}"), 400
            text = _extract_text_docx(docx_bytes)
            img_names = _extract_images_docx(docx_bytes, img_dir, base_id)
    except Exception as e:
        return jsonify(ok=False, message=f"Import failed: {e}"), 500

    _save_text_file(text, base_dir, base_id)  # optional

    base_url = f"/static/uploads/articles/{base_id}/img"
    img_urls = [f"{base_url}/{name}" for name in img_names]

    return jsonify(ok=True, text=text, images=img_urls)

def _extract_from_pdf(raw_bytes: bytes, img_dir: str, base_id: str):
    os.makedirs(img_dir, exist_ok=True)
    text_parts = []
    image_names = []

    doc = fitz.open(stream=raw_bytes, filetype="pdf")
    try:
        for page_index in range(len(doc)):
            page = doc[page_index]
            # text
            text_parts.append(page.get_text())

            # images
            for img in page.get_images(full=True):
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                try:
                    if pix.alpha:  # remove alpha for consistent PNGs
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    fname = f"p{page_index+1:03d}-{xref}.png"
                    out_path = os.path.join(img_dir, fname)
                    pix.save(out_path)
                    image_names.append(fname)
                finally:
                    pix = None
        full_text = "\n".join(text_parts).strip()
        return full_text, image_names
    finally:
        doc.close()


def _extract_text_docx(raw_bytes: bytes) -> str:
    doc = Document(io.BytesIO(raw_bytes))
    parts = [p.text for p in doc.paragraphs if p.text]
    # also consider tables (optional)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    text = "\n".join(parts)
    # normalize a bit
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text


def _extract_images_docx(raw_bytes: bytes, img_dir: str, base_id: str):
    os.makedirs(img_dir, exist_ok=True)
    image_names = []
    with zipfile.ZipFile(io.BytesIO(raw_bytes)) as zf:
        for name in zf.namelist():
            # embedded images live under word/media/
            if name.startswith("word/media/") and not name.endswith("/"):
                data = zf.read(name)
                fname = os.path.basename(name)
                out_path = os.path.join(img_dir, fname)
                with open(out_path, "wb") as f:
                    f.write(data)
                image_names.append(fname)
    return image_names

def _convert_doc_to_docx(raw_bytes: bytes) -> bytes:
    import tempfile, win32com.client as win32
    with tempfile.TemporaryDirectory() as td:
        src = os.path.join(td, "in.doc")
        dst = os.path.join(td, "out.docx")
        with open(src, "wb") as f: f.write(raw_bytes)

        word = win32.gencache.EnsureDispatch('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(src)
        try:
            wdFormatXMLDocument = 12
            doc.SaveAs(dst, FileFormat=wdFormatXMLDocument)
        finally:
            doc.Close(False)
            word.Quit()

        with open(dst, "rb") as f:
            return f.read()

def _save_text_file(text: str, base_dir: str, base_id: str):
    os.makedirs(base_dir, exist_ok=True)
    out_path = os.path.join(base_dir, "extracted.txt")
    try:
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(text or "")
    except Exception:
        # Non-fatal; ignore disk errors here
        pass

# (MW) ------AI: suggest title, clean up text, pick category------
@app.post("/subscriber/api/ai-suggest")
@login_required("Author", "Subscriber")
def api_ai_suggest():
    # Read JSON
    try:
        data = request.get_json(force=True) or {}
    except Exception:
        return jsonify(ok=False, message="Invalid JSON"), 400

    text = (data.get("text") or "").strip()
    if not text:
        return jsonify(ok=False, message="No text provided"), 400

    # Fetch categories via MySQL
    conn = get_db_connection()
    cur  = get_cursor(conn)
    cur.execute("SELECT categoryid, name FROM categories ORDER BY name")
    categories = cur.fetchall()
    cur.close(); conn.close()
    if not categories:
        return jsonify(ok=False, message="No categories configured"), 500

    # Heuristic title & cleanup
    def simple_title(t: str) -> str:
        for line in t.splitlines():
            s = line.strip()
            if len(s) >= 8:
                return s[:90].rstrip(" .!?:;")
        t = t.strip()
        return (t[:80] + "…") if len(t) > 80 else (t or "Untitled")

    def basic_cleanup(t: str) -> str:
        t = re.sub(r"[ \t]+", " ", t or "")
        t = re.sub(r"\n{3,}", "\n\n", t)
        return t.strip()

    title = simple_title(text)
    corrected = basic_cleanup(text)

    # Pick category by fuzzy match across full text
    choices = [c["name"] for c in categories]
    STOP = set("""
    a an the and or of for to in on with as at by from this that those these is are was were be been being
    it its they them their we our you your i me my he she his her which who whom whose will would should could
    """.split())

    words = re.findall(r"[A-Za-z]{3,}", corrected.lower())
    words = [w for w in words if w not in STOP and len(w) >= 4]

    from collections import Counter
    top_terms = [w for w, _n in Counter(words).most_common(30)] or ["general"]
    query = " ".join(top_terms)

    match = process.extractOne(query, choices, scorer=fuzz.WRatio)
    if match:
        cat_name, score, idx = match
        cat_id = categories[idx]["categoryid"]
    else:
        cat_id  = categories[0]["categoryid"]
        cat_name= categories[0]["name"]

    return jsonify(
        ok=True,
        title=title,
        corrected_text=corrected,
        category_id=cat_id,
        category_name=cat_name
    )

# (MW)
@app.route("/subscriber/api/articles", endpoint="subscriber_search_api")
@login_required("Author", "Subscriber")
def subscriber_search_api():
    cat    = request.args.get("cat")               # category name or ""
    q      = (request.args.get("q") or "").strip()
    limit  = int(request.args.get("limit", 30))
    offset = int(request.args.get("offset", 0))

    conn = get_db_connection()
    cur  = get_cursor(conn)

    where   = ["a.draft = 0", "a.visible = 1"]     # published & visible only
    params  = []
    joincat = "LEFT JOIN categories c ON a.catid = c.categoryid"

    if cat:
        where.append("c.name = %s")
        params.append(cat)

    if q:
        where.append("(a.title LIKE %s OR a.content LIKE %s)")
        like = f"%{q}%"
        params.extend([like, like])

    # Author-first showing in homepage 
    sql = f"""
      SELECT
        a.articleid, a.title, a.content, a.author,
        a.published_at, a.updated_at, a.image,
        c.name AS category
      FROM articles a
      {joincat}
      WHERE {" AND ".join(where)}
      ORDER BY
        CASE WHEN (EXISTS (
          SELECT 1 FROM users u
          WHERE TRIM(LOWER(u.name)) = TRIM(LOWER(a.author))
            AND TRIM(LOWER(u.usertype)) = 'author'
        )) THEN 0 ELSE 1 END,
        COALESCE(a.published_at, a.updated_at) DESC,
        a.articleid DESC
      LIMIT %s OFFSET %s
    """
    params.extend([limit, offset])

    cur.execute(sql, tuple(params))
    rows = cur.fetchall()
    cur.close(); conn.close()
    return jsonify(rows)

# (MW)
# -------View route (full-page article, shows pin state)----------------
@app.route("/subscriber/article/<int:article_id>")
@login_required("Author", "Subscriber")
def subscriber_article_view(article_id):
    conn = get_db_connection()
    cur  = get_cursor(conn)

    cur.execute("""
        SELECT
        a.articleid, a.title, a.content, a.author,
        a.published_at, a.updated_at, a.likes, a.dislikes,
        CASE
            WHEN a.image IS NULL OR a.image = '' THEN NULL
            WHEN a.image LIKE 'http%%' THEN a.image
            WHEN a.image LIKE '/static/%%' THEN a.image
            WHEN a.image LIKE '../static/%%' THEN REPLACE(a.image, '../static', '/static')
            ELSE CONCAT('/static/img/', a.image)
        END AS image,
        c.name AS category,
        u.userid        AS author_userid,
        u.usertype      AS author_usertype,  
        CASE
            WHEN u.image IS NULL OR u.image = '' THEN NULL
            WHEN u.image LIKE 'http%%' THEN u.image
            WHEN u.image LIKE '/static/%%' THEN u.image
            WHEN u.image LIKE '../static/%%' THEN REPLACE(u.image, '../static', '/static')
            ELSE CONCAT('/static/img/', u.image)
        END AS author_image
        FROM articles a
        LEFT JOIN categories c ON a.catid = c.categoryid
        LEFT JOIN users u      ON u.name = a.author
        WHERE a.articleid = %s
        AND (a.draft IS NULL OR a.draft = 0)
        LIMIT 1;
    """, (article_id,))

    article = cur.fetchone()
    if not article:
        cur.close(); conn.close()
        abort(404)

    # --- Views: +1 per session per article ---
    viewed_key = "viewed_articles"
    already_viewed = set(int(x) for x in session.get(viewed_key, []))
    if article_id not in already_viewed:
        cur.execute("UPDATE articles SET views = views + 1 WHERE articleid = %s", (article_id,))
        conn.commit()
        already_viewed.add(article_id)
        session[viewed_key] = list(already_viewed)

    # Pin state
    cur.execute(
        "SELECT 1 FROM subscriber_pins WHERE userid=%s AND articleid=%s LIMIT 1",
        (session["userid"], article_id)
    )
    is_pinned = cur.fetchone() is not None
    cur.close(); conn.close()

    # --- AI summarizer (reuse the same helper you use on public page) ---
    summary_points = []
    try:
        full_text = (article.get("content") or "").strip()
        summary_points = summarize_to_points(full_text)  # same helper as /article/<id>
    except Exception as e:
        summary_points = []

    return render_template(
        "subscriberArticleView.html",
        article=article,
        is_pinned=is_pinned,
        summary_points=summary_points
    )

@app.template_filter("dmyhm")
def fmt_ddmmyyyy_hhmm_ampm(dt):
    if not dt: return ""
    from datetime import datetime
    if isinstance(dt, str):
        try:
            dt = datetime.strptime(dt, "%Y-%m-%d %H:%M:%S")
        except ValueError:
            return dt
    s = dt.strftime("%d/%m/%Y %I:%M %p")
    return s[:-2] + s[-2:].lower()

def initials_from(name):
    parts = (name or "").strip().split()
    if not parts:
        return "U"
    return (parts[0][:1] + (parts[1][:1] if len(parts) > 1 else "")).upper()


@app.route("/subscriber/api/pin", methods=["POST"])
@login_required("Author", "Subscriber")
def subscriber_toggle_pin():
    data = request.get_json(silent=True) or {}
    article_id = data.get("articleid")

    if not article_id:
        return jsonify(ok=False, message="articleid is required"), 400

    conn = get_db_connection()
    cur  = get_cursor(conn)  # should be RealDictCursor

    try:
        # Check current state
        cur.execute("""
            SELECT id FROM subscriber_pins
            WHERE userid=%s AND articleid=%s
            LIMIT 1
        """, (session["userid"], article_id))
        row = cur.fetchone()

        if row:
            cur.execute("DELETE FROM subscriber_pins WHERE id=%s", (row['id'],))
            conn.commit()
            return jsonify(ok=True, state="unpinned", message="Article unpinned")
        else:
            cur.execute("""
                INSERT INTO subscriber_pins (userid, articleid)
                VALUES (%s, %s)
            """, (session["userid"], article_id))
            conn.commit()
            return jsonify(ok=True, state="pinned", message="Article pinned")
    except Exception:
        conn.rollback()
        return jsonify(ok=False, message="Unable to pin"), 500
    finally:
        cur.close()
        conn.close()

# ---------- Article Like/Dislike API ---------- 
@app.route('/subscriber/api/article/react', methods=['POST'])
@login_required("Author", "Subscriber")
def subscriber_article_react():
    data = request.get_json(silent=True) or {}
    try:
        articleid = int(data.get('articleid') or 0)
    except Exception:
        articleid = 0
    action = (data.get('action') or '').strip().lower()  # 'like' | 'dislike'

    if action not in ('like','dislike') or not articleid:
        return jsonify(ok=False, message='Bad request'), 400

    conn = get_db_connection()
    cur = get_cursor(conn)
    try:
        uid = session.get('userid')
        if uid is None:
            uname = (session.get('user') or '').strip()
            if not uname:
                return jsonify(ok=False, message='Not signed in'), 401
            cur_lookup = conn.cursor()
            try:
                cur_lookup.execute(
                    "SELECT userid FROM users WHERE username=%s OR name=%s LIMIT 1",
                    (uname, uname)
                )
                row = cur_lookup.fetchone()
            finally:
                cur_lookup.close()
            if not row:
                return jsonify(ok=False, message='User not found'), 401
            uid = int(row[0])
            session['userid'] = uid
        else:
            uid = int(uid)

        conn.start_transaction()

        cur.execute("""
            SELECT reaction
            FROM article_reactions
            WHERE articleid=%s AND userid=%s
            FOR UPDATE
        """, (articleid, uid))
        row = cur.fetchone()
        prev = (row['reaction'] if row else None)

        # compute deltas
        d_like = d_dislike = 0
        if prev is None:
            # new reaction
            if action == 'like':    d_like = 1
            if action == 'dislike': d_dislike = 1
            cur.execute("""
                INSERT INTO article_reactions (articleid, userid, reaction)
                VALUES (%s, %s, %s)
            """, (articleid, uid, action))
            state = action
        elif prev == action:
            # toggle off
            if action == 'like':    d_like = -1
            if action == 'dislike': d_dislike = -1
            cur.execute("""
                DELETE FROM article_reactions
                WHERE articleid=%s AND userid=%s
            """, (articleid, uid))
            state = 'none'
        else:
            # switch
            if prev == 'like' and action == 'dislike':
                d_like, d_dislike = -1, +1
            elif prev == 'dislike' and action == 'like':
                d_like, d_dislike = +1, -1
            cur.execute("""
                UPDATE article_reactions
                SET reaction=%s, updated_at=NOW()
                WHERE articleid=%s AND userid=%s
            """, (action, articleid, uid))
            state = action

        cur.execute("""
            UPDATE articles
            SET likes    = GREATEST(likes + %s, 0),
                dislikes = GREATEST(dislikes + %s, 0)
            WHERE articleid=%s
        """, (d_like, d_dislike, articleid))

        # return the fresh totals from articles
        cur.execute("SELECT likes, dislikes FROM articles WHERE articleid=%s FOR UPDATE", (articleid,))
        agg = cur.fetchone() or {'likes': 0, 'dislikes': 0}
        likes    = int(agg['likes'] or 0)
        dislikes = int(agg['dislikes'] or 0)

        conn.commit()
        return jsonify(ok=True, state=state, likes=likes, dislikes=dislikes)
    except Exception as e:
        conn.rollback()
        return jsonify(ok=False, message=f"Server error: {e}"), 500
    finally:
        try: cur.close()
        except: pass
        try: conn.close()
        except: pass

# --- My Articles: page ---
@app.route("/subscriber/my-articles")
@login_required("Author", "Subscriber")
def subscriber_my_articles():
    return render_template("subscriberMyArticles.html")

def _norm_image_path(p):
    if not p:
        return None
    p = p.strip()
    if p.startswith(("http://", "https://")):
        return p
    if p.startswith("/static/"):
        return p
    if p.startswith("./static/"):
        return p[1:]  
    if p.startswith(("uploads/", "img/")):
        return "/static/" + p              
    return "/static/uploads/" + p

# --- My Articles: JSON API ---
@app.route("/subscriber/api/my-articles", methods=["GET"])
@login_required("Subscriber")
def subscriber_api_my_articles():
    import psycopg2
    from psycopg2 import sql, extras

    user = session.get("user")
    username = user["username"] if isinstance(user, dict) else user

    category = request.args.get("category")
    status = request.args.get("status")
    search = request.args.get("search")
    limit = int(request.args.get("limit", 10))
    offset = int(request.args.get("offset", 0))

    # --- BASE QUERY ---
    query = """
        SELECT a.articleid AS article_id,
               a.title,
               a.published_at,
               a.updated_at,
               a.status,
               c.name AS category_name,
               u.name AS author_name
        FROM articles a
        JOIN categories c ON a.catid = c.categoryid
        JOIN users u ON a.author = u.name
        WHERE u.name = %s
    """
    params = [username]

    # --- FILTERS ---
    if category:
        query += " AND c.name = %s"
        params.append(category)

    if status:
        query += " AND a.status = %s"
        params.append(status)

    if search:
        query += " AND a.title ILIKE %s"
        params.append(f"%{search}%")

    # --- ORDER + PAGINATION (add only once) ---
    query += " ORDER BY a.published_at DESC LIMIT %s OFFSET %s"
    params.extend([limit, offset])

    # --- EXECUTE ---
    conn = get_db_connection()
    cur = get_cursor(conn)


    cur = conn.cursor(cursor_factory=extras.RealDictCursor)

    try:
        print("SQL:", query)
        print("PARAMS:", params)
        cur.execute(query, tuple(params))
        rows = cur.fetchall()
        cur.close()
        return jsonify({"success": True, "articles": rows})
    except Exception as e:
        import traceback
        print("ERROR:", traceback.format_exc())
        cur.close()
        return jsonify({"success": False, "error": str(e)}), 500



# View: load edit page 
@app.route("/subscriber/edit-article/<int:article_id>")
@login_required("Author", "Subscriber")
def subscriber_edit_article(article_id):
    conn = get_db_connection()
    cur  = get_cursor(conn)
    cur.execute("""
        SELECT a.articleid, a.title, a.content, a.catid, a.image, a.draft,
               a.published_at, a.updated_at
        FROM articles a
        WHERE a.articleid = %s AND a.author = %s
        LIMIT 1
    """, (article_id, session.get("user")))
    article = cur.fetchone()
    cur.close(); conn.close()
    if not article:
        return ("Forbidden", 403)
    return render_template("subscriberEditArticles.html", article=article)

# API: update article 
@app.route("/subscriber/api/article/<int:article_id>", methods=["POST"])
@login_required("Author", "Subscriber")
def subscriber_update_article(article_id):
    action      = (request.form.get("action") or "").lower()   # 'draft' | 'publish' | 'review'
    title       = (request.form.get("title") or "").strip()
    content     = (request.form.get("content") or "").strip()
    category_id = (request.form.get("category_id") or "").strip()
    image_file  = request.files.get("image")

    # Basic validation
    if not title or not content or not category_id:
        return jsonify(ok=False, message="Title, content and category are required."), 400

    conn = get_db_connection()
    cur  = get_cursor(conn)

    # Check ownership + fetch current status/visible/draft
    cur.execute("""
        SELECT author, status, visible, draft
        FROM articles
        WHERE articleid = %s
        LIMIT 1
    """, (article_id,))
    row = cur.fetchone()
    if not row or row["author"] != session.get("user"):
        cur.close(); conn.close()
        return jsonify(ok=False, message="Forbidden"), 403

    current_status = (row["status"] or "").lower()
    is_restricted  = current_status in ("reported", "pending_revision", "pending_approval")

    # Decide behavior by action
    is_draft   = (action == "draft")
    is_publish = (action == "publish")
    is_review  = (action == "review")

    # Enforce rule: restricted articles can ONLY be saved for review
    if is_restricted and not is_review:
        cur.close(); conn.close()
        return jsonify(ok=False, message="This article is under review. Only 'Save & Submit' is allowed."), 400

    # Handle image upload (optional)
    image_name = None
    if image_file and image_file.filename:
        image_name = image_file.filename
        image_path = os.path.join(app.static_folder, "img", image_name)
        image_file.save(image_path)

    # Build UPDATE fields common to all actions
    fields = ["title=%s","content=%s","catid=%s","updated_at=NOW()"]
    params = [title, content, category_id]

    flash_msg = "Saved."

    if is_review:
        # Your requirement:
        # status = pending_approval, draft = 1, visible = 0 (no publish_at)
        fields += ["status='pending_approval'", "draft=1", "visible=0"]
        flash_msg = "Submitted for review."
    elif is_draft:
        # Regular draft save (keep visible as-is)
        fields += ["draft=1"]
        flash_msg = "Draft saved."
    elif is_publish:
        # Publishing path (run moderation)
        # -> Profanity check
        profanity_check = check_profanity(title, content)
        if not profanity_check.get("ok", False):
            cur.close(); conn.close()
            return jsonify(ok=False, message=profanity_check.get("message","Failed profanity check.")), 400

        # -> AI moderation decides final visibility & message
        publish, visible, mod_msg = moderate_article(cur, title, content)
        # publish True means pass moderation; write publish timestamp
        if publish:
            fields += ["draft=0", "status='published'", "visible=%s", "published_at=NOW()"]
            params.append(int(bool(visible)))
        else:
            # If moderation says not publishable now, keep it as draft but reflect visibility
            fields += ["draft=1", "status='draft'", "visible=%s"]
            params.append(0)
        flash_msg = mod_msg or ("Published." if publish else "Saved (not published).")

    # Image if present
    if image_name:
        fields.append("image=%s")
        params.append(image_name)

    # WHERE + commit
    params.append(article_id)
    cur.execute(f"UPDATE articles SET {', '.join(fields)} WHERE articleid=%s", params)
    conn.commit()
    cur.close(); conn.close()
    return jsonify(ok=True, message=flash_msg, redirect=url_for("subscriber_my_articles"))

# Edit Reported Articles Page
@app.route("/subscriber/edit-reported-article/<int:article_id>")
@login_required("Author", "Subscriber")
def subscriber_edit_reported_article(article_id):
    conn = get_db_connection()
    cur = get_cursor(conn)
    cur.execute("""
        SELECT a.articleid, a.title, a.content, a.catid, a.image, a.draft,
               a.published_at, a.updated_at, a.status, a.visible
        FROM articles a
        WHERE a.articleid = %s AND a.author = %s
        LIMIT 1
    """, (article_id, session.get("user")))
    article = cur.fetchone()
    cur.close(); conn.close()
    if not article:
        return ("Forbidden", 403)

    # Guard: only allow reported/pending here (tweak if you like)
    if article["status"] not in ("reported", "pending_revision", "pending_approval"):
        return redirect(url_for("subscriber_edit_reported_article", article_id=article_id))

    # Reuse existing template; tell it we're in review mode
    return render_template("subscriberEditReportedArticles.html", article=article, review_mode=True)

# Bookmarks page
@app.route("/subscriber/bookmarks")
@login_required("Author", "Subscriber")
def subscriber_bookmarks():
    return render_template("subscriberBookmarks.html")

# Bookmarks data for the logged-in user
@app.route("/subscriber/api/bookmarks")
@login_required("Author", "Subscriber")
def subscriber_api_bookmarks():
    uid = session.get("userid")
    if uid is None:
        return jsonify(ok=False, message="Not logged in"), 401

    conn = get_db_connection()
    cur  = get_cursor(conn)
    cur.execute("""
        SELECT a.articleid, a.title, a.content, a.author,
               a.published_at, a.updated_at,
               CASE
                 WHEN a.image IS NULL OR a.image = '' THEN NULL
                 WHEN a.image LIKE 'http%%' THEN a.image
                 WHEN a.image LIKE '/static/%%' THEN a.image
                 WHEN a.image LIKE '../static/%%' THEN REPLACE(a.image, '../static', '/static')
                 ELSE CONCAT('/static/img/', a.image)
               END AS image,
               c.name AS category,
               sp.pinned_at
        FROM subscriber_pins sp
        JOIN articles a        ON a.articleid = sp.articleid
        LEFT JOIN categories c ON a.catid = c.categoryid
        WHERE sp.userid = %s
          AND (a.draft IS NULL OR a.draft = 0)
        ORDER BY sp.pinned_at DESC
    """, (uid,))
    rows = cur.fetchall()
    cur.close(); conn.close()
    return jsonify(rows)

# --- (MW) Subscriber comment functions ------

def _timeago(ts):
    if not ts:
        return ""
    if isinstance(ts, str):
        try:
            ts = datetime.fromisoformat(ts.replace("Z",""))
        except Exception:
            return ""
    now = datetime.now(timezone.utc)
    if ts.tzinfo is None:
        ts = ts.replace(tzinfo=timezone.utc)
    delta = int((now - ts).total_seconds())
    if delta < 60:   return "just now"
    if delta < 3600: return f"{delta//60}m ago"
    if delta < 86400:return f"{delta//3600}h ago"
    return f"{delta//86400}d ago"

def _to_iso(dt):
    try:    return dt.isoformat()
    except: return None

def _norm_img(p: Optional[str]) -> Optional[str]:
    if not p:
        return None
    if p.startswith("http"):
        return p
    if p.startswith("/static/"):
        return p
    if p.startswith("../static/"):
        return p.replace("../static", "/static")
    if p.startswith("./static/"):
        return p.replace("./static", "/static")
    return "/static/profilePictures/" + p.lstrip("/")

@app.get("/subscriber/api/comments")
@login_required("Author", "Subscriber")
def subscriber_api_comments():
    from app import get_db_connection  # if your factory exposes it here

    article_id = int(request.args.get("article_id", 0))
    uid = session.get("userid")  # your session key

    conn = get_db_connection()
    cur = get_cursor(conn)

    # 1) comments + author name + author image
    cur.execute("""
        SELECT
            c.commentid,
            c.comment_text,
            c.likes,
            c.dislikes,
            c.created_at,
            c.userid,
            c.is_reply,
            c.reply_to_comment_id,
            u.name  AS author,
            u.image AS author_image,
            u.usertype      AS author_usertype,
            (c.userid = %s) AS mine
        FROM comments c
        JOIN users u ON u.userid = c.userid
        WHERE c.articleid = %s AND c.visible = 1
        ORDER BY c.created_at DESC
    """, (uid, article_id))
    rows = cur.fetchall()

    # 2) my reactions for these comments
    my_reaction_map = {}
    if rows:
        ids = [r["commentid"] for r in rows]
        placeholders = ",".join(["%s"] * len(ids))
        cur.execute(f"""
            SELECT commentid, reaction
            FROM comment_reactions
            WHERE userid = %s AND commentid IN ({placeholders})
        """, (uid, *ids))
        for r in cur.fetchall():
            my_reaction_map[r["commentid"]] = r["reaction"]

    cur.close(); conn.close()

    out = []
    for r in rows:
        out.append({
            "commentid": r["commentid"],
            "text": r["comment_text"],
            "likes": int(r["likes"] or 0),
            "dislikes": int(r["dislikes"] or 0),
            "created_at": _to_iso(r["created_at"]),
            "author": r["author"] or "Anonymous",
            "author_image": _norm_img(r.get("author_image")),
            "mine": bool(r["mine"]),
            "my_reaction": my_reaction_map.get(r["commentid"]),
            "is_reply": bool(r.get("is_reply")),
            "parent_id": r.get("reply_to_comment_id"),
            "author_id": r["userid"],
            "author_usertype": (r.get("author_usertype") or "")
        })
    return jsonify(out)

@app.post("/subscriber/api/comments")
@login_required("Author", "Subscriber")
def subscriber_api_comment_create():
    from app import get_db_connection

    uid = session.get("userid")
    article_id = request.form.get("articleid")
    text = request.form.get("text", "").strip()
    parent_id = request.form.get("parent_id")

    if not text:
        return jsonify({"ok": False, "error": "Empty comment."}), 400

    conn = get_db_connection()
    cur = get_cursor(conn)

    # Ensure 1-level replies only (reply to top-level)
    top_parent_id = None
    if parent_id:
        try:
            pid = int(parent_id)
            cur.execute("SELECT reply_to_comment_id FROM comments WHERE commentid=%s", (pid,))
            row = cur.fetchone()
            if row:
                top_parent_id = row["reply_to_comment_id"] or pid
            else:
                top_parent_id = pid
        except Exception:
            top_parent_id = None

    cur.close()
    cur = get_cursor(conn)
    cur.execute("""
    INSERT INTO comments (articleid, userid, comment_text, is_reply, reply_to_comment_id)
    VALUES (%s, %s, %s, %s, %s)
    """, (
        article_id,
        uid,
        text,
        1 if top_parent_id else 0,  # convert boolean to smallint
        top_parent_id  # None is fine for NULL
    ))

    conn.commit()
    cur.close(); conn.close()
    return jsonify({"ok": True})

@app.put("/subscriber/api/comments/<int:comment_id>")
@login_required("Author", "Subscriber")
def subscriber_api_comment_update(comment_id):
    from app import get_db_connection

    uid = session.get("userid")
    text = (request.json.get("text") or "").strip()
    if not text:
        return jsonify(ok=False, message="text required"), 400

    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("SELECT userid FROM comments WHERE commentid=%s", (comment_id,))
    row = cur.fetchone()
    if not row:
        cur.close(); conn.close(); return jsonify(ok=False, message="Not found"), 404
    if row[0] != uid:
        cur.close(); conn.close(); return jsonify(ok=False, message="Forbidden"), 403

    cur.execute("UPDATE comments SET comment_text=%s WHERE commentid=%s", (text, comment_id))
    conn.commit()
    cur.close(); conn.close()
    return jsonify(ok=True)

@app.delete("/subscriber/api/comments/<int:comment_id>")
@login_required("Author", "Subscriber")
def subscriber_api_comment_delete(comment_id):
    from app import get_db_connection

    uid = session.get("userid")
    conn = get_db_connection()
    cur = conn.cursor()
    cur.execute("SELECT userid FROM comments WHERE commentid=%s", (comment_id,))
    row = cur.fetchone()
    if not row:
        cur.close(); conn.close(); return jsonify(ok=False, message="Not found"), 404
    if row[0] != uid:
        cur.close(); conn.close(); return jsonify(ok=False, message="Forbidden"), 403
    cur.execute("DELETE FROM comments WHERE commentid=%s", (comment_id,))
    conn.commit()
    cur.close(); conn.close()
    return jsonify(ok=True)

@app.post("/subscriber/api/comments/<int:comment_id>/react")
@login_required("Author", "Subscriber")
def subscriber_api_comment_react(comment_id):
    from app import get_db_connection

    uid = session.get("userid")
    action = (request.json.get("action") or "").lower()  # "like"|"dislike"|"clear"
    if action not in ("like", "dislike", "clear"):
        return jsonify(ok=False, message="Invalid action"), 400

    conn = get_db_connection()
    cur = get_cursor(conn)

    # Fetch previous reaction by this user
    cur.execute(
        "SELECT reaction FROM comment_reactions WHERE commentid=%s AND userid=%s",
        (comment_id, uid)
    )
    row = cur.fetchone()
    prev = row[0] if row else None

    # Fetch current likes/dislikes from the comment
    cur.execute("SELECT likes, dislikes FROM comments WHERE commentid=%s", (comment_id,))
    cd = cur.fetchone()
    if not cd:
        cur.close()
        conn.close()
        return jsonify(ok=False, message="Comment not found"), 404

    # ✅ Convert to integers to prevent TypeError
    likes, dislikes = int(cd[0] or 0), int(cd[1] or 0)

    if action == "clear":
        if prev == "like":
            likes = max(0, likes - 1)
        elif prev == "dislike":
            dislikes = max(0, dislikes - 1)
        if prev:
            cur.execute(
                "DELETE FROM comment_reactions WHERE commentid=%s AND userid=%s",
                (comment_id, uid)
            )
    else:
        if prev == action:
            # Clicking the same reaction again clears it
            if action == "like":
                likes = max(0, likes - 1)
            else:
                dislikes = max(0, dislikes - 1)
            cur.execute(
                "DELETE FROM comment_reactions WHERE commentid=%s AND userid=%s",
                (comment_id, uid)
            )
            action = "clear"
        else:
            # Remove previous reaction if exists
            if prev == "like":
                likes = max(0, likes - 1)
            elif prev == "dislike":
                dislikes = max(0, dislikes - 1)

            # Apply new reaction
            if action == "like":
                likes += 1
            else:
                dislikes += 1

            # ✅ PostgreSQL-safe upsert
            cur.execute("""
                INSERT INTO comment_reactions (commentid, userid, reaction, updated_at)
                VALUES (%s, %s, %s, CURRENT_TIMESTAMP)
                ON CONFLICT (commentid, userid) DO UPDATE
                  SET reaction = EXCLUDED.reaction,
                      updated_at = CURRENT_TIMESTAMP
            """, (comment_id, uid, action))

    # Update counts in the comments table
    cur.execute(
        "UPDATE comments SET likes=%s, dislikes=%s WHERE commentid=%s",
        (likes, dislikes, comment_id)
    )

    conn.commit()
    cur.close()
    conn.close()

    return jsonify(
        ok=True,
        likes=likes,
        dislikes=dislikes,
        state=None if action == "clear" else action
    )


@app.route("/subscriber/report_article", methods=["POST"])
@login_required("Author", "Subscriber")
def report_article():
    data = request.get_json(silent=True) or {}
    # frontend sends "id"; keep backward compat with "articleID"
    article_id = data.get("id") or data.get("articleid")
    reason     = (data.get("reason") or "").strip()
    details    = (data.get("details") or "").strip()

    reporter_id   = session.get("userid")
    reporter_name = (session.get("user") or "").strip()

    if not reporter_id:
        return jsonify(ok=False, message="User not logged in."), 401
    if not article_id or not reason:
        return jsonify(ok=False, message="Article ID and reason are required."), 400

    conn = get_db_connection()
    cur  = get_cursor(conn)
    try:
        # 1) verify article exists + author name
        cur.execute("SELECT author FROM articles WHERE articleid = %s", (article_id,))
        row = cur.fetchone()
        if not row:
            return jsonify(ok=False, message="Article not found."), 404

        # 2) block self-report (by author name, case-insensitive)
        art_author = (row.get("author") or "").strip()
        if art_author and reporter_name and art_author.lower() == reporter_name.lower():
            return jsonify(ok=False, message="You cannot report your own article."), 403

        # 3) de-dupe (same user reporting same article)
        cur.execute(
            "SELECT 1 FROM article_reports WHERE article_id = %s AND reporter_id = %s LIMIT 1",
            (article_id, reporter_id)
        )
        if cur.fetchone():
            return jsonify(ok=False, message="You have already reported this article."), 400

        # 4) insert
        cur.execute("""
            INSERT INTO article_reports (article_id, reporter_id, reason, details, created_at)
            VALUES (%s, %s, %s, %s, NOW())
        """, (article_id, reporter_id, reason, details))
        cur.execute("""
            UPDATE articles
            SET status = 'reported', visible = 1, updated_at = NOW()
            WHERE articleid = %s AND status IN ('draft','published')
        """, (article_id,))
        conn.commit()
        return jsonify(ok=True, message="Report submitted successfully.")
    except mysql.connector.Error as err:
        conn.rollback()
        print("DB error:", err)
        return jsonify(ok=False, message="Internal error while submitting report."), 500
    finally:
        cur.close(); conn.close()

@app.route("/subscriber/report_comment", methods=["POST"])
@login_required("Author", "Subscriber")
def report_comment():
    data = request.get_json(silent=True) or {}
    comment_id = data.get("id") or data.get("commentid")
    reason     = (data.get("reason") or "").strip()
    details    = (data.get("details") or "").strip()

    reporter_id = session.get("userid")
    if not reporter_id:
        return jsonify(ok=False, message="User not logged in."), 401
    if not comment_id or not reason:
        return jsonify(ok=False, message="Comment ID and reason are required."), 400

    conn = get_db_connection()
    cur  = get_cursor(conn)
    try:
        cur.execute(
            "SELECT 1 FROM comment_reports WHERE comment_id = %s AND reporter_id = %s LIMIT 1",
            (comment_id, reporter_id)
        )
        if cur.fetchone():
            return jsonify(ok=False, message="You have already reported this comment."), 400

        cur.execute("""
            INSERT INTO comment_reports (comment_id, reporter_id, reason, details, created_at)
            VALUES (%s, %s, %s, %s, NOW())
        """, (comment_id, reporter_id, reason, details))
        conn.commit()
        return jsonify(ok=True, message="Report submitted successfully.")
    except mysql.connector.Error as err:
        conn.rollback()
        print("DB error:", err)
        return jsonify(ok=False, message="Internal error while submitting report."), 500
    finally:
        cur.close(); conn.close()

# (MW) --------------Profile page----------------
def save_profile_image(file_storage):
    if not file_storage or not file_storage.filename:
        return None
    safe = secure_filename(file_storage.filename)
    name, ext = os.path.splitext(safe)
    if ext.lower() not in {".jpg", ".jpeg", ".png", ".webp", ".gif"}:
        raise ValueError("Unsupported image type")
    unique = f"{name}_{int(time.time())}{ext}"
    dest_dir = Path(app.root_path) / "static" / "profilePictures"
    dest_dir.mkdir(parents=True, exist_ok=True)
    file_storage.save(dest_dir / unique)
    return f"/static/profilePictures/{unique}"

@app.route("/profile")
@login_required("Author", "Subscriber")
def profile():
    uid = session["userid"]

    conn = get_db_connection()
    cur  = get_cursor(conn)
    cur.execute("""
        SELECT userid, name, email, usertype, image, bio
        FROM users
        WHERE userid=%s
        LIMIT 1
    """, (uid,))
    u = cur.fetchone()
    cur.close(); conn.close()

    if not u:
        flash("User not found.", "error")
        return redirect(url_for("subscriberHomepage"))

    img = (u.get("image") or "")
    if img.startswith("../static/"):
        u["image"] = img.replace("../static", "/static")
    elif img.startswith("./static/"):
        u["image"] = img.replace("./static", "/static")

    profile = {
        "display_name": u["name"],
        "email": u["email"],
        "usertype": u["usertype"],
        "bio": u.get("bio"),
        "avatar_url": u.get("image"),
    }

    return render_template(
        "subscriberProfile.html",
        user=u,
        profile=profile,
        pinned=[],
        articles=[]
    )

@app.route("/profile/update", methods=["POST"])
@login_required("Author", "Subscriber")
def profile_update():
    uid = session["userid"]
    next_url = request.form.get("next") or request.referrer or url_for("subscriberHomepage")

    name = (request.form.get("name") or "").strip()
    bio  = (request.form.get("bio")  or "").strip()
    avatar_file = request.files.get("avatar")

    conn = get_db_connection()
    cur  = get_cursor(conn)
    cur.execute("SELECT name, bio, image FROM users WHERE userid=%s LIMIT 1", (uid,))
    row = cur.fetchone()
    cur.close()

    current_name = (row.get("name") or "").strip() if row else ""
    current_bio  = (row.get("bio") or "").strip() if row else ""

    no_text_change = (name == current_name and bio == current_bio)
    no_avatar_upload = not (avatar_file and avatar_file.filename)

    if no_text_change and no_avatar_upload:
        return redirect(url_for("profile", nochange=1, next=next_url))

    new_img = None
    if avatar_file and avatar_file.filename:
        new_img = save_profile_image(avatar_file)

    try:
        cur = get_cursor(conn)
        if new_img:
            cur.execute("UPDATE users SET name=%s, bio=%s, image=%s WHERE userid=%s",
                        (name, bio, new_img, uid))
        else:
            cur.execute("UPDATE users SET name=%s, bio=%s WHERE userid=%s",
                        (name, bio, uid))
        conn.commit()
    except Exception as e:
        conn.rollback()
        err = str(e)
        return redirect(url_for("profile", error=1, msg=err, next=next_url))
    finally:
        cur.close(); conn.close()

    return redirect(url_for("profile", updated=1, next=next_url))

from flask import jsonify

# only Authors can downgrade to Subscriber
@app.post("/api/profile/downgrade-author")
@login_required("Author")  
def api_downgrade_author():
    uid   = session["userid"]
    role  = session.get("usertype") or session.get("role") or "Author"

    if role != "Author":
        return jsonify(ok=False, message="Only Author accounts can downgrade."), 403

    conn = get_db_connection()
    try:
        cur = conn.cursor()
        cur.execute("""
            UPDATE users
               SET previous_usertype = usertype,
                   usertype = 'Subscriber'
             WHERE userid = %s
             LIMIT 1
        """, (uid,))
        conn.commit()
    except Exception as e:
        conn.rollback()
        return jsonify(ok=False, message=str(e)), 500
    finally:
        cur.close(); conn.close()

    session["usertype"] = "Subscriber"
    session["role"] = "Subscriber"

    return jsonify(ok=True,
                   message="Role downgraded to Subscriber.",
                   redirect=url_for("subscriberHomepage"))

# only Subscribers can cancel sunscription
@app.post("/api/profile/cancel-subscription")
@login_required("Subscriber")  
def api_cancel_subscription():
    uid  = session["userid"]
    role = session.get("usertype") or session.get("role") or "Subscriber"

    if role != "Subscriber":
        return jsonify(ok=False, message="Only Subscriber accounts can be cancelled here."), 403

    conn = get_db_connection()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM users WHERE userid=%s LIMIT 1", (uid,))
        conn.commit()
    except Exception as e:
        conn.rollback()
        return jsonify(ok=False, message=str(e)), 500
    finally:
        cur.close(); conn.close()

    session.clear()

    return jsonify(ok=True,
                   message="Subscription cancelled.",
                   redirect=url_for("index"))

@app.route("/subscriber/profile/<int:user_id>")
@login_required("Author", "Subscriber")
def subscriber_profile_view(user_id):
    conn = get_db_connection()
    cur  = get_cursor(conn)

    cur.execute("""
        SELECT userid, name, email, usertype, image, bio
        FROM users
        WHERE userid=%s
    """, (user_id,))
    user = cur.fetchone()
    if not user:
        cur.close(); conn.close()
        return ("User not found", 404)

    img = (user.get("image") or "").strip()
    show_img = None
    if img:
        if img.startswith("http://") or img.startswith("https://"):
            show_img = img
        elif img.startswith("/static/"):
            show_img = img
        elif img.startswith("../static/"):
            show_img = img.replace("../static", "/static")
        elif img.startswith("./static/"):
            show_img = img.replace("./static", "/static")
        else:
            show_img = f"/static/profilePictures/{img}"
    user["image_url"] = show_img

    cur.execute("""
        SELECT a.articleid, a.title, a.published_at, a.image
        FROM articles a
        WHERE a.author = %s
          AND a.visible = 1
          AND a.status = 'published'
        ORDER BY COALESCE(a.published_at, a.updated_at) DESC, a.articleid DESC
        LIMIT 12
    """, (user["name"],))
    articles = cur.fetchall()

    cur.close(); conn.close()
    return render_template("subscriberViewProfile.html", user=user, articles=articles)

@app.get("/subscriber/api/profile/<int:user_id>/articles")
@login_required("Author", "Subscriber")
def api_profile_articles(user_id):
    conn = get_db_connection()
    cur = get_cursor(conn)

    cur.execute("SELECT name FROM users WHERE userid=%s", (user_id,))
    u = cur.fetchone()
    if not u:
        cur.close(); conn.close()
        return jsonify(ok=False, message="User not found"), 404

    limit  = int(request.args.get("limit", 24))
    offset = int(request.args.get("offset", 0))

    cur.execute("""
        SELECT a.articleid, a.title, a.published_at,
               CASE
                 WHEN a.image IS NULL OR a.image = '' THEN NULL
                 WHEN a.image LIKE 'http%%' THEN a.image
                 WHEN a.image LIKE '/static/%%' THEN a.image
                 WHEN a.image LIKE '../static/%%' THEN REPLACE(a.image,'../static','/static')
                 WHEN a.image LIKE './static/%%'  THEN REPLACE(a.image,'./static','/static')
                 ELSE CONCAT('/static/img/', a.image)
               END AS image
        FROM articles a
        WHERE a.author = %s
          AND a.visible = 1
          AND a.status  = 'published'
        ORDER BY COALESCE(a.published_at, a.updated_at) DESC, a.articleid DESC
        LIMIT %s OFFSET %s
    """, (u["name"], limit, offset))
    rows = cur.fetchall()
    cur.close(); conn.close()
    return jsonify(ok=True, items=rows)

@app.get("/subscriber/author/<int:author_id>")
@login_required("Author", "Subscriber")
def subscriber_article_list_by_author(author_id):
    conn = get_db_connection(); 
    cur = get_cursor(conn)
    cur.execute("SELECT userid, name, image, bio FROM users WHERE userid=%s", (author_id,))
    author = cur.fetchone()

    cur.execute("""
        SELECT articleid, title, content, image, published_at, updated_at
        FROM articles
        WHERE author = %s AND draft = 0
        ORDER BY COALESCE(published_at, updated_at) DESC
    """, (author_id,))
    articles = cur.fetchall()
    cur.close(); conn.close()

    return render_template("subscriberAuthorArticles.html",
                           author=author, articles=articles)

@app.post("/subscriber/api/articles/<int:article_id>/delete")
@login_required("Author", "Subscriber")
def subscriber_delete_article(article_id):
    username = session.get("user")  # current logged-in subscriber
    conn = get_db_connection()
    cur = get_cursor(conn)

    # Make sure the article belongs to the logged-in user
    cur.execute("SELECT * FROM articles WHERE articleid = %s AND author = %s", (article_id, username))
    article = cur.fetchone()

    if not article:
        cur.close(); conn.close()
        return jsonify({"error": "Article not found or access denied"}), 403

    try:
        cur.execute("DELETE FROM articles WHERE articleid = %s", (article_id,))
        conn.commit()
        cur.close(); conn.close()
        return jsonify({"success": True, "deleted_id": article_id})
    except Exception as e:
        conn.rollback()
        cur.close(); conn.close()
        return jsonify({"error": str(e)}), 500
    
# (MW) --- Ads API: random visible ad for sidebar rotation ---
@app.route("/api/ads/random")
def api_random_ad():
    conn = get_db_connection()
    cur  = get_cursor(conn)
    # Random 1 where visible=1
    cur.execute("""
        SELECT adsid, adstitle, adsdescription, adsimage, adswebsite
        FROM advertisement
        WHERE visible = 1
        ORDER BY RANDOM()
        LIMIT 1
    """)
    row = cur.fetchone()
    cur.close(); conn.close()

    if not row:
        return jsonify({"ok": False, "message": "No ads available"}), 404

    def norm_image(p: str) -> str:
        if not p: return "/static/img/placeholder.png"
        if p.startswith("http"): return p
        if p.startswith("/static/"): return p
        if p.startswith("../static/"): return p.replace("../static", "/static")
        if p.startswith("./static/"): return p.replace("./static", "/static")
        return "/static/ads/" + p.lstrip("/")

    return jsonify({
        "ok": True,
        "ad": {
            "id":        row["adsid"],
            "title":     row["adstitle"],
            "desc":      row.get("adsdescription") or "",
            "image":     norm_image(row["adsimage"]),
            "website":   row["adswebsite"]
        }
    })
    
# (MW) ---------- SubscriberAnalytics ----------
@app.route("/subscriber/analytics")
@login_required("Author", "Subscriber")
def subscriber_analytics():
    return render_template("subscriberAnalytics.html")

@app.route("/subscriber/api/analytics/my-articles")
@login_required("Author", "Subscriber")
def subscriber_analytics_my_articles():
    author_name = session.get("user", "")

    conn = get_db_connection()
    cur  = get_cursor(conn)

    cur.execute("""
        SELECT
            a.articleid,
            a.title,
            a.published_at,
            COALESCE(a.views, 0)          AS views,
            COALESCE(bm.bookmarks, 0)     AS bookmarks,
            COALESCE(cc.comment_count, 0) AS comment_count,
            COALESCE(a.likes, 0)          AS likes,
            COALESCE(a.dislikes, 0)       AS dislikes
        FROM articles a
        LEFT JOIN (
            SELECT articleid, COUNT(*) AS bookmarks
            FROM subscriber_pins
            GROUP BY articleid
        ) bm ON bm.articleid = a.articleid
        LEFT JOIN (
            SELECT articleid, COUNT(*) AS comment_count
            FROM comments
            GROUP BY articleid
        ) cc ON cc.articleid = a.articleid
        WHERE a.draft = 0
          AND a.visible = 1
          AND a.author = %s
        ORDER BY a.published_at DESC
    """, (author_name,))

    rows = cur.fetchall() or []

    from datetime import datetime
    for r in rows:
        dt = r.get("published_at")
        if hasattr(dt, "strftime"):
            r["published_at"] = f'{dt.strftime("%a")}, {dt.strftime("%d-%m-%Y %I:%M")} {dt.strftime("%p").lower()}'
        elif isinstance(dt, str) and dt:
            for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%dT%H:%M:%S", "%a, %d %b %Y %H:%M:%S %Z"):
                try:
                    d = datetime.strptime(dt, fmt)
                    r["published_at"] = f'{d.strftime("%a")}, {d.strftime("%d-%m-%Y %I:%M")} {d.strftime("%p").lower()}'
                    break
                except ValueError:
                    pass
        r["shares"] = 0

    cur.close()
    conn.close()
    return jsonify(rows)

def _row_donation(idx, row):
    donation_id     = row["donation_id"]
    donation_amount = row["donation_amount"] 
    payment_method  = row["payment_method"] or ""
    ts              = row["paymentdatetime"]  

    if isinstance(ts, (datetime, )):
        payment_date = ts.strftime("%d-%m-%Y")
        payment_time = ts.strftime("%I:%M %p").lower() 
    else:
        payment_date = str(ts)
        payment_time = ""

    return {
        "no": idx,

        "donation_id": donation_id,
        "donation_amount": float(donation_amount or 0),
        "payment_method": payment_method,
        "paymentdatetime": ts,

        "payment_date": payment_date,      
        "payment_time": payment_time,      
        "donation_id_display": donation_id 
    }

@app.get("/subscriber/api/analytics/my-donations")
@login_required("Author", "Subscriber")
def subscriber_analytics_my_donations():
    uid = session.get("userid")
    if not uid:
        return jsonify({"error": "unauthorized"}), 401

    conn = get_db_connection()
    cur  = get_cursor(conn)
    try:
        cur.execute("""
            SELECT
                d.donation_id,
                d.donation_amount,
                d.payment_method,
                d.paymentdatetime,
                di.card_brand
            FROM donations AS d
            LEFT JOIN donation_info AS di
              ON di.donation_id = d.donation_id
            WHERE d.userid = %s
            ORDER BY d.paymentDateTime DESC, d.donation_id DESC
        """, (uid,))
        recs = cur.fetchall() or []
    finally:
        cur.close()
        conn.close()

    def _shape(row):
        ts = row["paymentdatetime"]
        if hasattr(ts, "strftime"):
            payment_date = ts.strftime("%d-%m-%Y")
            payment_time = ts.strftime("%I:%M %p").lower()
        else:
            payment_date, payment_time = "", ""

        pm = (row.get("payment_method") or "").lower()
        if pm == "card":
            pm_disp = row.get("card_brand") or "Card"
        elif pm == "paynow":
            pm_disp = "PayNow"
        else:
            pm_disp = pm.capitalize() if pm else "—"

        return {
            "donation_id": row["donation_id"],
            "payment_method": pm_disp,
            "donation_amount": float(row.get("donation_amount") or 0),
            "payment_date": payment_date,
            "payment_time": payment_time,
        }

    return jsonify([_shape(r) for r in recs])

@app.get("/subscriber/api/analytics/overview")
@login_required("Author", "Subscriber")
def subscriber_analytics_overview():
    user_id     = session.get("userid")
    author_name = session.get("user", "")

    conn = get_db_connection()
    cur  = get_cursor(conn)
    try:
        # Total views of user's articles
        cur.execute("""
            SELECT COALESCE(SUM(views), 0) AS total_views
            FROM articles
            WHERE author = %s AND draft = 0 AND visible = 1
        """, (author_name,))
        row = cur.fetchone()
        total_views = int(row['total_views'] or 0) if row else 0

        # Total bookmarks
        cur.execute("""
            SELECT COUNT(*) AS total_bookmarks
            FROM subscriber_pins
            WHERE userid = %s
        """, (user_id,))
        row = cur.fetchone()
        total_bookmarks = int(row['total_bookmarks'] or 0) if row else 0

        # Likes and dislikes by me
        try:
            cur.execute("""
                SELECT
                  COALESCE(SUM(CASE WHEN reaction='like' THEN 1 ELSE 0 END),0) AS likes_by_me,
                  COALESCE(SUM(CASE WHEN reaction='dislike' THEN 1 ELSE 0 END),0) AS dislikes_by_me
                FROM comment_reactions
                WHERE userid = %s
            """, (user_id,))
            row = cur.fetchone()
            likes_by_me    = int(row['likes_by_me'] or 0) if row else 0
            dislikes_by_me = int(row['dislikes_by_me'] or 0) if row else 0
        except Exception:
            likes_by_me, dislikes_by_me = 0, 0

        # Likes and dislikes on my articles
        cur.execute("""
            SELECT
              COALESCE(SUM(likes), 0)    AS my_article_likes,
              COALESCE(SUM(dislikes), 0) AS my_article_dislikes
            FROM articles
            WHERE author = %s AND draft = 0 AND visible = 1
        """, (author_name,))
        row = cur.fetchone()
        my_article_likes    = int(row['my_article_likes'] or 0) if row else 0
        my_article_dislikes = int(row['my_article_dislikes'] or 0) if row else 0

        # Contributed amount
        cur.execute("""
            SELECT COALESCE(SUM(donation_amount),0) AS contributed_amount
            FROM donations
            WHERE userid = %s
        """, (user_id,))
        row = cur.fetchone()
        contributed_amount = float(row['contributed_amount'] or 0.0) if row else 0.0

    finally:
        cur.close()
        conn.close()

    return jsonify({
        "total_views": total_views,
        "total_bookmarks": total_bookmarks,
        "comment_likes": likes_by_me,
        "comment_dislikes": dislikes_by_me,
        "contributed_amount": contributed_amount,
        "my_article_likes": my_article_likes,
        "my_article_dislikes": my_article_dislikes,
    })



# (MW)
# ---------- Donations-----------
@app.get("/donate")
@login_required("Author", "Subscriber")
def donate_form():
    return render_template("subscriberDonation.html")

@app.post("/donate")
@login_required("Author", "Subscriber")
def donate_submit():
    """
    1) Validate form
    2) INSERT into donations (master)
    3) INSERT into donation_info (detail)
    """

    method = (request.form.get("payment_method") or "").strip().lower()
    amount_raw = (request.form.get("donation_amount") or "").strip()

    card_brand = (request.form.get("card_brand") or "").strip() or None
    cardnumber = (request.form.get("cardnumber") or "").strip() or None
    paynowref  = (request.form.get("paynowref")  or "").strip() or None

    try:
        amt = decimal.Decimal(amount_raw)
    except Exception:
        flash("Invalid amount.", "danger")
        return redirect(url_for("donate_form"))

    if amt <= 0:
        flash("Amount must be greater than 0.", "danger")
        return redirect(url_for("donate_form"))

    if method not in {"card", "paynow"}:
        flash("Please select a valid payment method.", "danger")
        return redirect(url_for("donate_form"))

    if method == "card":
        if not cardnumber or len(cardnumber) != 16 or not cardnumber.isdigit():
            flash("Invalid card number.", "danger")
            return redirect(url_for("donate_form"))
        if not card_brand:
            card_brand = "Unknown"

    if method == "paynow":
        if not paynowref:
            flash("PayNow reference is required.", "danger")
            return redirect(url_for("donate_form"))

    user_id = session.get("userid")
    if not user_id:
        flash("You are not logged in.", "danger")
        return redirect(url_for("login"))

    conn = get_db_connection()
    cur  = get_cursor(conn)
    try:
        cur.execute(
            """
            INSERT INTO donations
              (userid, donation_amount, payment_method, paymentdatetime, created_by)
            VALUES
              (%s, %s, %s, NOW(), %s)
            """,
            (user_id, str(amt), method, user_id)
        )
        conn.commit()

        cur.execute("SELECT LAST_INSERT_ID()")
        donation_id = cur.fetchone()[0]

        if method == "card":
            cur.execute(
                """
                INSERT INTO donation_info
                  (donation_id, card_brand, cardnumber)
                VALUES
                  (%s, %s, %s)
                """,
                (donation_id, card_brand, cardnumber)
            )
        elif method == "paynow":
            cur.execute(
                """
                INSERT INTO donation_info
                  (donation_id, paynowref)
                VALUES
                  (%s, %s)
                """,
                (donation_id, paynowref)
            )

        conn.commit()

        return redirect(url_for("donate_form", thanks=1, next=url_for("subscriberHomepage")))

    except Exception as e:
        app.logger.exception("DONATION_INSERT_FAILED: %s", e)

        msg = str(e)
        if "foreign key constraint fails" in msg.lower():
            flash("Insert failed: your user account (userid) must exist in 'users' table.", "danger")
        else:
            flash(f"Could not save donation: {e}", "danger")
        conn.rollback()
        return redirect(url_for("donate_form"))
    finally:
        try:
            cur.close()
            conn.close()
        except Exception:
            pass

# (YY)
# ---------- Auth ----------
@app.route("/login", methods=["GET", "POST"])
def login():
    remembered_email = request.cookies.get("remembered_email")  # Get cookie if exists

    if request.method == "POST":
        email = request.form["email"]
        password = request.form["password"]
        remember = request.form.get("remember")

        conn = get_db_connection()
        cursor = get_cursor(conn)
        cursor.execute("SELECT * FROM users WHERE email = %s", (email,))
        user = cursor.fetchone()

        if user:
            # Replace this if you store plain text passwords (not recommended)
            # if password == user["password"]:
            if password == user["password"]:  # or use check_password_hash(user["password"], password)
                # Check if suspended
                if user["usertype"].lower() == "suspended":
                    flash("Your account has been suspended. Please contact support.")
                    return redirect(url_for("login"))

                # ✅ Mark user as logged in and update last active time
                cursor.execute("""
                    UPDATE users
                    SET is_logged_in = 1, last_active = NOW()
                    WHERE userid = %s
                """, (user["userid"],))

                # ✅ Record login activity (optional but useful)
                cursor.execute("""
                    INSERT INTO login_activity (userid, email, login_time, ip_address)
                    VALUES (%s, %s, NOW(), %s)
                """, (user["userid"], user["email"], request.remote_addr))

                conn.commit()

                # ✅ Set session data
                session["userid"] = user["userid"]
                session["usertype"] = user["usertype"]
                session["user"] = user["name"]

                role_redirects = {
                    "Admin": "adminHomepage",
                    "Moderator": "modHomepage",
                    "Subscriber": "subscriberHomepage",
                    "Author": "subscriberHomepage"
                }

                redirect_route = role_redirects.get(user["usertype"])
                resp = make_response(redirect(url_for(redirect_route) if redirect_route else url_for("login")))

                # ✅ Handle "Remember Me" cookie
                if remember:
                    resp.set_cookie("remembered_email", email, max_age=30 * 24 * 60 * 60)
                else:
                    resp.delete_cookie("remembered_email")

                if not redirect_route:
                    flash("Invalid user type.")
                return resp

            else:
                flash("Incorrect password.")
                return redirect(url_for("login"))
        else:
            flash("User not found.")
            return redirect(url_for("login"))

    # On GET, render login page with remembered email if any
    return render_template("login.html", remembered_email=remembered_email)

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        name = request.form["name"]
        email = request.form["email"]
        password = request.form["password"]

        conn = get_db_connection()
        cursor = get_cursor(conn)
        try:
            cursor.execute(
                "INSERT INTO users (name, email, password, usertype) VALUES (%s, %s, %s, %s)",
                (name, email, password, "Subscriber")
            )
            conn.commit()
            return redirect(url_for("login"))
        except mysql.connector.IntegrityError:
            # Instead of returning text, render template with error
            return render_template("register.html", error="This email is already registered. Please use a different email.")
        finally:
            cursor.close()
            conn.close()
    return render_template("register.html")

@app.route("/logout")
def logout():
    # Update database to mark user inactive before clearing session
    if "userid" in session:
        user_id = session["userid"]
        conn = get_db_connection()
        cursor = get_cursor(conn)
        cursor.execute("""
            UPDATE users 
            SET is_logged_in = 0, last_active = NOW()
            WHERE userid = %s
        """, (user_id,))
        conn.commit()
        cursor.close()
        conn.close()

    # Destroy session and clear cookie
    session.clear()
    resp = redirect(url_for("login"))
    resp.delete_cookie("session")  # remove Flask session cookie
    flash("You have been logged out successfully.")
    return resp

@app.before_request
def update_last_active():
    if "userid" in session:
        conn = get_db_connection()
        cursor = get_cursor(conn)
        cursor.execute("""
            UPDATE users
            SET last_active = NOW()
            WHERE userid = %s
        """, (session["userid"],))
        conn.commit()
        cursor.close()
        conn.close()

@app.after_request
def add_no_cache_headers(response):
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, post-check=0, pre-check=0, max-age=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "-1"
    return response

# ---------- Moderator: manage users ----------
@app.route("/manageUsers")
@login_required("Moderator")
def manage_users():
    if "userid" not in session or session.get("usertype") != "Moderator":
        flash("Access denied.")
        return redirect(url_for("login"))

    conn = get_db_connection()
    cursor = get_cursor(conn)
    cursor.execute("""
        SELECT userid, name, email, usertype
        FROM users
        WHERE usertype IN ('Subscriber', 'Author', 'Suspended')
    """)
    users = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template("manageUsers.html", users=users)

@app.route("/warnUser", methods=["POST"])
@login_required("Moderator")
def warn_user():
    if "userid" not in session or session.get("usertype") != "Moderator":
        flash("Access denied.")
        return redirect(url_for("login"))

    warned_user_id = request.form.get("userid")

    conn = get_db_connection()
    cursor = get_cursor(conn)

    warning_message = "You have received a warning from the moderator."
    cursor.execute(
        "INSERT INTO warnings (userid, message) VALUES (%s, %s)",
        (warned_user_id, warning_message)
    )
    conn.commit()
    cursor.close()
    conn.close()

    flash("Warning sent successfully!")
    return redirect(url_for("manage_users"))

@app.route("/toggleSuspend", methods=["POST"])
@login_required("Moderator")
def toggle_suspend():
    if "userid" not in session or session.get("usertype") != "Moderator":
        flash("Access denied.")
        return redirect(url_for("login"))

    user_id = request.form["userid"]

    conn = get_db_connection()
    cursor = get_cursor(conn)

    cursor.execute("SELECT name, usertype, previous_usertype FROM users WHERE userid = %s", (user_id,))
    user = cursor.fetchone()

    if user:
        if user["usertype"] == "Suspended":
            restored_type = user["previous_usertype"] if user["previous_usertype"] else "Subscriber"
            cursor.execute("""
                UPDATE users
                SET usertype = %s, previous_usertype = NULL
                WHERE userid = %s
            """, (restored_type, user_id))
            flash(f"{user['name']} has been unsuspended.")
        else:
            cursor.execute("""
                UPDATE users
                SET previous_usertype = usertype, usertype = 'Suspended'
                WHERE userid = %s
            """, (user_id,))
            flash(f"{user['name']} has been suspended.")

    conn.commit()
    cursor.close()
    conn.close()

    return redirect(url_for("manage_users"))

# ---------- Forgot password ----------
@app.route("/forgot-password", methods=["GET", "POST"])
def forgot_password():
    if request.method == "POST":
        email = request.form["email"]
        new_password = request.form["new_password"]
        confirm_password = request.form["confirm_password"]

        if new_password != confirm_password:
            flash("Passwords do not match.")
            return redirect(url_for("forgot_password"))

        conn = get_db_connection()
        cursor = get_cursor(conn)
        cursor.execute("SELECT * FROM users WHERE email = %s", (email,))
        user = cursor.fetchone()

        if user:
            cursor.execute(
                "UPDATE users SET password = %s WHERE email = %s",
                (new_password, email)
            )
            conn.commit()
            flash("Password has been reset. Please login with the new password.")
            cursor.close()
            conn.close()
            return redirect(url_for("login"))
        else:
            flash("Email not found.")
            cursor.close()
            conn.close()
            return redirect(url_for("forgot_password"))
    return render_template("forgot_password.html")

# ---------- Change Password ----------
@app.post("/api/profile/change-password")
@login_required("Author", "Subscriber", "Moderator", "Admin")
def api_change_password():
    uid = session.get("userid")
    if not uid:
        return jsonify(ok=False, message="Not signed in."), 401

    data = request.get_json(silent=True) or {}
    old = (data.get("old") or "").strip()
    new = (data.get("new") or "").strip()
    confirm = (data.get("confirm") or "").strip()

    # client-side should validate too, but keep server-side guardrails
    if not old or not new or not confirm:
        return jsonify(ok=False, message="All fields are required."), 400
    if new != confirm:
        return jsonify(ok=False, message="New passwords do not match."), 400
    if new == old:
        return jsonify(ok=False, message="New password must be different."), 400
    if len(new) < 4:
        return jsonify(ok=False, message="Password must be at least 4 characters."), 400

    conn = get_db_connection()
    cur  = conn.cursor()
    # Atomic check+update. This only succeeds when the old password matches THIS user.
    cur.execute(
        "UPDATE users SET password=%s WHERE userid=%s AND password=%s",
        (new, uid, old)
    )
    conn.commit()
    changed = cur.rowcount
    cur.close(); conn.close()

    if changed != 1:
        # No rows updated → old password was wrong
        return jsonify(ok=False, message="Current password is incorrect."), 400

    return jsonify(ok=True, message="Password updated.")

# ---------- Admin views ----------
@app.route("/viewAllUsers")
@login_required("Admin")
def view_all_users():
    if "userid" not in session or session.get("usertype") != "Admin":
        flash("Access denied.")
        return redirect(url_for("login"))

    conn = get_db_connection()
    cursor = get_cursor(conn)
    cursor.execute("SELECT userid, name, email, usertype FROM users")
    users = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template("viewAllUsers.html", users=users)

@app.route("/searchAccount", methods=["GET", "POST"])
@login_required("Admin")
def search_account():
    if "userid" not in session or session.get("usertype") != "Admin":
        flash("Access denied.")
        return redirect(url_for("login"))

    users = []
    search_term = ""

    if request.method == "POST":
        search_term = request.form["search_term"]
        conn = get_db_connection()
        cursor = get_cursor(conn)
        cursor.execute("""
            SELECT userid, name, email, usertype
            FROM users
            WHERE name LIKE %s OR email LIKE %s
        """, (f"%{search_term}%", f"%{search_term}%"))
        users = cursor.fetchall()
        cursor.close()
        conn.close()

    return render_template("searchAccount.html", users=users, search_term=search_term)

@app.route("/createUser", methods=["GET", "POST"])
@login_required("Admin")
def create_user():
    if "userid" not in session or session.get("usertype") != "Admin":
        flash("Access denied.")
        return redirect(url_for("login"))

    if request.method == "POST":
        name = request.form["name"]
        email = request.form["email"]
        password = request.form["password"]
        confirm_password = request.form["confirm_password"]
        usertype = request.form["usertype"]

        if password != confirm_password:
            flash("Passwords do not match.")
            return redirect(url_for("create_user"))

        conn = get_db_connection()
        cursor = get_cursor(conn)
        cursor.execute(
            "INSERT INTO users (name, email, password, usertype) VALUES (%s, %s, %s, %s)",
            (name, email, password, usertype)
        )
        conn.commit()
        cursor.close()
        conn.close()

        flash("User created successfully!")
        return redirect(url_for("create_user"))

    return render_template("createUser.html")

@app.route("/updateUser", methods=["GET", "POST"])
@login_required("Admin")
def update_user():
    if "userid" not in session or session.get("usertype") != "Admin":
        flash("Access denied.")
        return redirect(url_for("login"))

    conn = get_db_connection()
    cursor = get_cursor(conn)

    if request.method == "POST":
        email = request.form["email"]

        # Check if user exists and is a Subscriber
        cursor.execute("SELECT * FROM users WHERE email = %s", (email,))
        user = cursor.fetchone()

        if not user:
            flash("User not found.")
        elif user["usertype"] != "Subscriber":
            flash("Only Subscribers can be updated to Author.")
        else:
            # Change Subscriber → Author
            cursor.execute(
                "UPDATE users SET usertype = 'Author' WHERE email = %s",
                (email,),
            )
            conn.commit()
            flash("User successfully updated to Author.")

    cursor.close()
    conn.close()

    return render_template("updateUser.html")

@app.route("/manageUserStatus", methods=["GET", "POST"])
@login_required("Admin")
def manage_user_status():
    if "userid" not in session or session.get("usertype") != "Admin":
        flash("Access denied.")
        return redirect(url_for("login"))

    conn = get_db_connection()
    cursor = get_cursor(conn)

    if request.method == "POST":
        user_id = request.form.get("userid")
        action = request.form.get("action")

        # Get current user info
        cursor.execute("SELECT usertype, previous_usertype FROM users WHERE userid = %s", (user_id,))
        user = cursor.fetchone()

        if not user:
            flash("User not found.")
        else:
            if action == "suspend":
                if user["usertype"] == "Admin":
                    flash("You cannot suspend an Admin account.")
                elif user["usertype"] != "Suspended":
                    cursor.execute(
                        "UPDATE users SET previous_usertype = usertype, usertype = 'Suspended' WHERE userid = %s",
                        (user_id,)
                    )
                    flash("User suspended successfully.")

            elif action == "reactivate":
                if user["usertype"] == "Suspended" and user["previous_usertype"]:
                    cursor.execute(
                        "UPDATE users SET usertype = previous_usertype, previous_usertype = NULL WHERE userid = %s",
                        (user_id,)
                    )
                    flash("User reactivated successfully.")
                else:
                    flash("No previous role found. Cannot reactivate.")

            elif action == "delete":
                if user["usertype"] == "Admin":
                    flash("You cannot delete an Admin account.")
                else:
                    cursor.execute("DELETE FROM users WHERE userid = %s", (user_id,))
                    flash("User deleted successfully.")

        conn.commit()

    # Fetch all users
    cursor.execute("SELECT userid, name, email, usertype, previous_usertype FROM users")
    users = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template("manageUserStatus.html", users=users)

@app.route("/newUsers")
@login_required("Admin")
def report_new_users():
    if "userid" not in session or session.get("usertype") != "Admin":
        flash("Access denied.")
        return redirect(url_for("login"))

    conn = get_db_connection()
    cursor = get_cursor(conn)

    # Example: users created in the last 7 days
    cursor.execute("""
        SELECT userid, name, email, usertype, created_at
        FROM users
        WHERE created_at >= NOW() - INTERVAL '7 days'
        ORDER BY created_at DESC
    """)
    new_users = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template("newUsers.html", users=new_users)

@app.route("/articleSubmission")
@login_required("Admin")
def article_submission():
    # Ensure only Admins can access
    if "userid" not in session or session.get("usertype") != "Admin":
        flash("Access denied.")
        return redirect(url_for("login"))

    conn = get_db_connection()
    cursor = get_cursor(conn)

    # Fetch only articles created by valid users (exclude system/orphan articles)
    cursor.execute("""
    SELECT 
        a.articleid,
        a.title,
        a.content,
        a.author,
        a.published_at,
        a.updated_at,
        a.image,
        a.catid,
        a.draft
    FROM articles a
    INNER JOIN users u ON a.author = u.name
    WHERE u.usertype IN ('Author', 'Subscriber')
    ORDER BY a.published_at DESC
    """)

    articles = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template("articleSubmissions.html", articles=articles)

@app.route("/loginActivity")
@login_required("Admin")
def login_activity():
    if "userid" not in session or session.get("usertype") != "Admin":
        flash("Access denied.")
        return redirect(url_for("login"))

    conn = get_db_connection()
    cursor = get_cursor(conn)

    cursor.execute("""
        SELECT 
            userid, 
            email,
            is_logged_in,
            last_active
        FROM users
        ORDER BY last_active DESC
    """)
    users = cursor.fetchall()

    # Calculate "x hours ago"
    now = datetime.now()
    for user in users:
        if user["last_active"]:
            diff = now - user["last_active"]
            seconds = diff.total_seconds()
            if seconds < 60:
                user["last_seen"] = "Just Now"
            elif seconds < 3600:
                user["last_seen"] = f"{int(seconds // 60)} min ago"
            elif seconds < 86400:
                user["last_seen"] = f"{int(seconds // 3600)} hrs ago"
            else:
                user["last_seen"] = f"{int(seconds // 86400)} days ago"
        else:
            user["last_seen"] = "N/A"

    cursor.close()
    conn.close()

    return render_template("loginActivity.html", users=users)

@app.route("/flaggedArticles")
@login_required("Moderator")
def flagged_articles():
    conn = get_db_connection()
    cursor = get_cursor(conn)

    cursor.execute("""
        SELECT 
            ar.report_id,
            ar.article_id AS articleid,
            a.title,
            a.author,
            a.published_at AS article_created,
            ar.reason AS flagged_reason,
            ar.details AS flagged_details,
            ar.created_at AS flagged_at
        FROM article_reports ar
        LEFT JOIN articles a ON ar.article_id = a.articleid
        WHERE ar.status = 'pending'
        ORDER BY ar.created_at DESC
    """)
    articles = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template("flaggedArticles.html", articles=articles)

@app.route("/reviewArticle", methods=["POST"])
@login_required("Moderator")
def review_article():
    report_id = request.form.get("report_id")
    article_id = request.form.get("articleid")
    action = request.form.get("action")

    if action not in ("approve", "reject"):
        flash("Invalid request.", "error")
        return redirect("/flaggedArticles")

    conn = get_db_connection()
    cursor = get_cursor(conn)

    try:
        # === CASE 1: From flaggedArticles (has report_id) ===
        if report_id:
            cursor.execute("SELECT article_id FROM article_reports WHERE report_id = %s", (report_id,))
            report = cursor.fetchone()
            if not report:
                flash("Report not found.", "error")
                return redirect("/flaggedArticles")
            
            article_id = report["article_id"]

            if action == "approve":  # moderator agrees with report
                cursor.execute("UPDATE article_reports SET status = 'reviewed' WHERE article_id = %s", (article_id,))
                cursor.execute("UPDATE articles SET visible = 0, status = 'pending_revision' WHERE articleid = %s", (article_id,))
            else:  # moderator rejects the report
                cursor.execute("UPDATE article_reports SET status = 'dismissed' WHERE report_id = %s", (report_id,))
                cursor.execute("UPDATE articles SET visible = 1, status = 'published' WHERE articleid = %s", (article_id,))

            redirect_target = "/flaggedArticles"

        # === CASE 2: From pendingArticles (no report_id) ===
        elif article_id:
            if action == "approve":
                cursor.execute("UPDATE articles SET status = 'published', visible = 1 WHERE articleid = %s", (article_id,))
            elif action == "reject":
                cursor.execute("UPDATE articles SET status = 'pending_revision', visible = 0 WHERE articleid = %s", (article_id,))

            redirect_target = "/pendingArticles"

        else:
            flash("Invalid data provided.", "error")
            return redirect("/flaggedArticles")

        conn.commit()
        flash(f"Article {action}d successfully.", "success")

    except mysql.connector.Error as err:
        print("Database error:", err)
        flash("An error occurred while processing the request.", "error")

    finally:
        cursor.close()
        conn.close()

    return redirect(redirect_target)

@app.route("/getArticle/<int:article_id>")
def get_article(article_id):
    conn = get_db_connection()
    cursor = get_cursor(conn)
    cursor.execute("""
        SELECT title, content, author, published_at, image
        FROM articles
        WHERE articleid = %s
    """, (article_id,))
    article = cursor.fetchone()
    cursor.close()
    conn.close()

    if not article:
        return jsonify({"error": "Article not found"}), 404

    return jsonify(article)

# (YY)
@app.route("/flaggedComments")
@login_required("Moderator")
def flagged_comments():
    conn = get_db_connection()
    cursor = get_cursor(conn)

    try:
        cursor.execute("""
            SELECT 
                cr.report_id,
                cr.comment_id AS commentid,
                c.comment_text,
                u.name AS user,      -- join users to get name
                cr.reason AS flagged_reason,
                cr.details AS flagged_details,
                cr.created_at AS flagged_at
            FROM comment_reports cr
            LEFT JOIN comments c ON cr.comment_id = c.commentid
            LEFT JOIN users u ON c.userid = u.userid   -- get the comment author's name
            WHERE cr.status = 'pending'
            ORDER BY cr.created_at DESC
        """)
        comments = cursor.fetchall()

    except mysql.connector.Error as err:
        print("Database error:", err)
        comments = []

    finally:
        cursor.close()
        conn.close()

    return render_template("flaggedComments.html", comments=comments)

@app.route("/reviewComment", methods=["POST"])
@login_required("Moderator")
def reviewComment():
    comment_id = request.form.get("commentid")
    action = request.form.get("action")  # "approve" or "reject"

    if not comment_id or action not in ("approve", "reject"):
        flash("Invalid request.", "danger")
        return redirect("/flaggedComments")

    conn = get_db_connection()
    cursor = get_cursor(conn)

    try:
        # Get all reports for this comment
        cursor.execute("SELECT report_id FROM comment_reports WHERE comment_id = %s", (comment_id,))
        reports = cursor.fetchall()
        if not reports:
            flash("No reports found for this comment.", "danger")
            return redirect("/flaggedComments")

        if action == "approve":
            # Mark all reports as reviewed
            cursor.execute(
                "UPDATE comment_reports SET status = 'reviewed' WHERE comment_id = %s",
                (comment_id,)
            )
            # Hide the comment from public
            cursor.execute(
                "UPDATE comments SET visible = 0 WHERE commentid = %s",
                (comment_id,)
            )
        elif action == "reject":
            # Mark all reports as dismissed
            cursor.execute(
                "UPDATE comment_reports SET status = 'dismissed' WHERE comment_id = %s",
                (comment_id,)
            )
            # Ensure comment stays visible
            cursor.execute(
                "UPDATE comments SET visible = 1 WHERE commentid = %s",
                (comment_id,)
            )

        conn.commit()
        flash(f"Comment {action}d successfully.", "success")

    except mysql.connector.Error as err:
        print("Database error:", err)
        flash("An error occurred while reviewing the comment.", "danger")

    finally:
        cursor.close()
        conn.close()

    return redirect("/flaggedComments")

@app.route("/getComment/<int:comment_id>")
def get_comment(comment_id):
    conn = get_db_connection()
    cursor = get_cursor(conn)
    cursor.execute("""
        SELECT c.commentid, c.comment_text, c.created_at, u.name AS user
        FROM comments c
        JOIN users u ON c.userid = u.userid
        WHERE c.commentid = %s
    """, (comment_id,))
    comment = cursor.fetchone()
    cursor.close()
    conn.close()

    if not comment:
        return jsonify({"error": "Comment not found"}), 404

    return jsonify(comment)

# (YY)
@app.route("/pendingArticles")
@login_required("Moderator")
def pending_articles():
    conn = get_db_connection()
    cursor = get_cursor(conn)

    cursor.execute("""
        SELECT 
            articleid,
            title,
            author,
            updated_at
        FROM articles
        WHERE status = 'pending_approval'
        ORDER BY updated_at DESC
    """)
    articles = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template("pendingArticles.html", articles=articles)

@app.route("/manageCategories", methods=["GET"])
@login_required("Moderator")
def manage_categories():
    conn = get_db_connection()
    cursor = get_cursor(conn)
    cursor.execute("SELECT * FROM categories")
    categories = cursor.fetchall()
    cursor.close()
    conn.close()
    return render_template("manageCategories.html", categories=categories)
	
# (YY)
@app.route("/manageCategories/create", methods=["GET", "POST"])
@login_required("Moderator")
def add_category():
    if request.method == "POST":
        name = request.form.get("categoryName").strip()
        description = request.form.get("categoryDescription").strip()

        if not name:
            flash("Category name cannot be empty.", "error")
            return redirect(url_for("add_category"))

        conn = get_db_connection()
        cursor = get_cursor(conn)

        # Check if category already exists
        cursor.execute("SELECT * FROM categories WHERE name = %s", (name,))
        existing = cursor.fetchone()
        if existing:
            flash("Category already exists.", "error")
            cursor.close()
            conn.close()
            return redirect(url_for("add_category"))

        # Insert new category with description
        cursor.execute(
            "INSERT INTO categories (name, description) VALUES (%s, %s)",
            (name, description)
        )
        conn.commit()
        cursor.close()
        conn.close()

        flash("Category added successfully.", "success")
        return redirect(url_for("manage_categories"))

    # GET request
    return render_template("createCategory.html")
# (YY)
@app.route("/manageCategories/update", methods=["GET", "POST"])
@login_required("Moderator")
def update_category():
    conn = get_db_connection()
    cursor = get_cursor(conn)

    if request.method == "POST":
        categoryid = request.form.get("categoryid")
        name = request.form.get("categoryName").strip()
        description = request.form.get("categoryDescription").strip()

        if not name:
            flash("Category name cannot be empty.", "error")
            return redirect(url_for("update_category"))

        cursor.execute(
            "UPDATE categories SET name=%s, description=%s WHERE categoryid=%s",
            (name, description, categoryid)
        )
        conn.commit()
        flash("Category updated successfully.", "success")
        return redirect(url_for("manage_categories"))

    # GET request: load all categories to select which to update
    cursor.execute("SELECT * FROM categories")
    categories = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template("updateCategory.html", categories=categories)

# (YY)
@app.route("/manageCategories/delete", methods=["GET"])
@login_required("Moderator")
def delete_category_page():
    search = request.args.get("search", "")
    conn = get_db_connection()
    cursor = get_cursor(conn)

    if search:
        cursor.execute("SELECT * FROM categories WHERE name LIKE %s ORDER BY name ASC", ('%' + search + '%',))
    else:
        cursor.execute("SELECT * FROM categories ORDER BY name ASC")

    categories = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template("deleteCategory.html", categories=categories)

# Handle actual deletion (YY)
@app.route("/manageCategories/delete/<int:categoryid>", methods=["POST"])
@login_required("Moderator")
def delete_category(categoryid):
    conn = get_db_connection()
    cursor = get_cursor(conn)
    cursor.execute("DELETE FROM categories WHERE categoryid=%s", (categoryid,))
    conn.commit()
    cursor.close()
    conn.close()
    flash("Category deleted successfully.", "success")
    return redirect(url_for("delete_category_page"))


# ---------- AI Summarizer -------
# to summarize text and convert to bullet points
def summarize_to_points(text, sentences_count=5):
    if not text.strip():
        return ["No summary available"]

    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = LsaSummarizer()
    summary = summarizer(parser.document, sentences_count)
    return [str(s) for s in summary]

# endpoint: to receive article text and return bullet-point summary
@app.route("/summarize", methods=["POST"])
def summarize():
    data = request.get_json()
    text = data.get("text", "")

    try:
        bullet_points = summarize_to_points(text)
        return jsonify({"summary": bullet_points})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ---------- Multilingual SUpport ----------
translated = GoogleTranslator(source="auto", target="de").translate
@app.route("/translate", methods=["POST"])
def translate():
    try:
        data = request.get_json()
        title = data.get("title", "")
        content = data.get("content", "")
        summary = data.get("summary", [])
        target_lang = data.get("target_lang", "en")

        if target_lang == "en":
            return jsonify({
                "title": title,
                "content": content,
                "summary": summary
            })

        # translate title
        translated_title = GoogleTranslator(source="auto", target=target_lang).translate(title) if title else ""

        # split content into chunks of 2-3 sentences
        sentence_endings = re.compile(r'(?<=[.!?])\s+')
        sentences = sentence_endings.split(content.strip())
        
        chunks = []
        chunk_size = 3  # number of sentences per chunk
        for i in range(0, len(sentences), chunk_size):
            chunks.append(" ".join(sentences[i:i+chunk_size]))

        translated_paragraphs = []
        for chunk in chunks:
            try:
                translated_paragraphs.append(GoogleTranslator(source="auto", target=target_lang).translate(chunk))
            except Exception:
                translated_paragraphs.append(chunk)  # fallback

        translated_content = "\n\n".join(translated_paragraphs)

        # translate summary
        if summary:
            full_summary = "\n".join(summary)
            try:
                translated_full_summary = GoogleTranslator(source="auto", target=target_lang).translate(full_summary)
                translated_summary = translated_full_summary.split("\n")
            except Exception:
                translated_summary = summary  # fallback
        else:
            translated_summary = []

        return jsonify({
            "title": translated_title,
            "content": translated_content,
            "summary": translated_summary
        })

    except Exception as e:
        print("Error in /translate:", e)
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500
    
# ---------- chatbot support ----------
@app.route("/chat", methods=["POST"])
def chat():
    user_input = request.json.get("message", "")
    response = getChatbotResponse(user_input)
    return jsonify({"response": response})

# ---------- Dev helper ----------
def open_browser():
    webbrowser.open_new("http://127.0.0.1:5000/")

if __name__ == "__main__":
    threading.Timer(1.0, open_browser).start()
    app.run(debug=True)
	